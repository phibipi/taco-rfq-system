import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import time
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docxtpl import DocxTemplate
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn  
from docx.oxml import parse_xml, OxmlElement 
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
import urllib.parse
import io
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import hashlib
import requests


# --- MAIN APP ---
st.set_page_config(page_title="TACO Procurement", layout="wide", page_icon="🚛")

# --- SHEET CONNECTION ---
SPREADSHEET_ID = "1j9GCq8Wwm-MM8hOamsH26qlmjNwuDBuEMnbw6ORzTQk"


# --- UI PAGE FONT BUTTON ETC ---
def init_style():
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&family=Montserrat:wght@500;600;700&display=swap');
        
        /* GLOBAL TEXT */
        html, body, p, label, span, div, .stMarkdown, .stText, h1, h2, h3, h4, h5, h6, th, td {
            font-family: 'Inter', sans-serif;
            color: #111827 !important;
        }
        .stApp { background-color: #F3F4F6; }
        
        /* HEADER */
        h1, h2, h3, h4 { font-family: 'Montserrat', sans-serif !important; font-weight: 600 !important; }
        
        /* CARD */
        div[data-testid="stVerticalBlockBorderWrapper"] {
            background-color: #FFFFFF; border-radius: 10px; border: 1px solid #E5E7EB; padding: 18px; box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        }
        
        /* INPUTS */
        .stTextInput input, .stSelectbox div[data-baseweb="select"] > div, .stNumberInput input, textarea {
            background-color: #F9FAFB !important; border: 1px solid #D1D5DB !important; border-radius: 6px; color: #111827 !important;
        }
        
        /* DROPDOWN & MULTISELECT */
        div[data-baseweb="popover"], div[data-baseweb="menu"], ul[data-baseweb="menu"] { background-color: #FFFFFF !important; }
        li[role="option"] { background-color: #FFFFFF !important; color: #111827 !important; }
        li[role="option"]:hover { background-color: #F3F4F6 !important; }
        span[data-baseweb="tag"] { background-color: #E5E7EB !important; color: #111827 !important; }
        
        /* --- STANDARD BUTTONS --- */
        div[data-testid="stButton"] button {
            background-color: #FCA568 !important; color: #FFFFFF !important; border: 2px solid #F58536 !important;
            border-radius: 8px !important; font-weight: 600 !important; padding: 0.5rem 1.2rem !important;
            box-shadow: 0 2px 4px rgba(249, 115, 22, 0.2);
        }
        div[data-testid="stButton"] button:hover {
            background-color: #EA580C !important; transform: translateY(-1px);
        }
        button[kind="secondary"] { background-color: #E5E7EB !important; color: #111827 !important; box-shadow: none !important; }
        button[kind="secondary"]:hover { background-color: #D1D5DB !important; }
        
        /* --- TOMBOL SIMPAN (SUBMIT) LEBIH BESAR --- */
        div[data-testid="stFormSubmitButton"] button {
            background-color: #FCA568 !important; 
            color: #FFFFFF !important; 
            border: 2px solid #F58536 !important;
            border-radius: 12px !important; 
            font-weight: 900 !important;      /* Lebih Tebal */
            font-size: 24px !important;        /* Huruf Lebih Besar */
            padding: 0.8rem 2rem !important;  /* Ukuran Tombol Lebih Besar */
            box-shadow: 0 4px 8px rgba(249, 115, 22, 0.3);
            width: 100% !important;            /* Tombol Memanjang Penuh */
            transition: all 0.3s ease;
        }
        div[data-testid="stFormSubmitButton"] button:hover {
            background-color: #EA580C !important; 
            transform: translateY(-2px);
            box-shadow: 0 6px 12px rgba(249, 115, 22, 0.4);
        }
        
        /* --- DISABLED BUTTON STYLE (ABU TUA) --- */
        div[data-testid="stButton"] button:disabled {
            background-color: #6B7280 !important; /* Cool Gray 500 (Abu Tua) */
            color: #FFFFFF !important;
            border: 1px solid #4B5563 !important;
            cursor: not-allowed !important;
            opacity: 1 !important; /* Memaksa warna keluar jelas */
            box-shadow: none !important;
        }
        
        /* TABLE FIX */
        div[data-testid="stDataFrame"] { background-color: #F8FAFC !important; border-radius: 8px; border: 1px solid #E5E7EB; }
        div[data-testid="stDataEditor"] input { color: #000000 !important; background-color: #FFFFFF !important; caret-color: #000000 !important; -webkit-text-fill-color: #000000 !important; }
        
        /* STATUS */
        .status-done { color: #065F46 !important; background-color: #D1FAE5; padding: 5px 12px; border-radius: 6px; font-size: 14px; font-weight: 700; }
        .status-pending { color: #B91C1C !important; background-color: #FEE2E2 !important; padding: 6px 16px; border-radius: 8px; font-size: 15px !important; font-weight: 800 !important; border: 1px solid #FECACA; display: inline-block; }
        .route-dest-list { font-size: 13px; color: #4B5563 !important; line-height: 1.4; }
        .streamlit-expanderHeader { background-color: #FFFFFF; border: 1px solid #E5E7EB; border-radius: 8px; color: #111827 !important; font-weight: 600; }
        
        hr { margin-top: 0.5em; margin-bottom: 0.5em; border: none; height: 1px; background-color: #E5E7EB; }
        
        /* TABS */
        .stTabs [data-baseweb="tab-list"] { gap: 8px; }
        .stTabs [data-baseweb="tab"] { background-color: #FFFFFF; border: 1px solid #E5E7EB; border-radius: 6px; padding: 4px 16px; color: #111827 !important; }
        .stTabs [data-baseweb="tab"][aria-selected="true"] { background-color: #B5B2B0 !important; color: #FFFFFF !important; border: none; }
       
        /* --- PERBAIKAN TAMPILAN EXPANDER (DROPDOWN) --- */
        
        /* 1. Target Kotak Utama Expander */
        div[data-testid="stExpander"] {
            background-color: #FFFFFF !important;
            border: 1px solid #E5E7EB !important;
            border-radius: 8px !important;
            color: #111827 !important; /* Warna teks isi gelap */
            box-shadow: 0 1px 2px rgba(0,0,0,0.05);
        }

        /* 2. Target Bagian Detail (Isi Dalam) */
        div[data-testid="stExpander"] details {
            background-color: #FFFFFF !important;
            border-radius: 8px !important;
        }

        /* 3. Target Judul (Header) Expander */
        div[data-testid="stExpander"] summary {
            background-color: #FFFFFF !important; /* Paksa Background Putih */
            color: #111827 !important;             /* Paksa Teks Hitam */
            font-weight: 600 !important;
            border-radius: 8px !important;
        }
        
        /* 4. Target Ikon Panah Kecil (Agar tidak putih/hilang) */
        div[data-testid="stExpander"] summary svg {
            fill: #6B7280 !important;  /* Warna panah abu tua */
            color: #6B7280 !important;
        }

        /* 5. Efek Hover (Opsional: Garis jadi Orange) */
        div[data-testid="stExpander"]:hover {
            border-color: #FCA568 !important;
        }
        
        </style>
    """, unsafe_allow_html=True)


# --- KONEKSI & CACHE (SMART DETECTION) ---
@st.cache_resource
def connect_to_gsheet():
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    
    try:
        # 1. CEK APAKAH DI STREAMLIT CLOUD (SECRETS)
        if "gcp_service_account" in st.secrets:
            
            creds_dict = st.secrets["gcp_service_account"]
            creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
            client = gspread.authorize(creds)
            return client.open_by_key(SPREADSHEET_ID)

        # 2. CEK APAKAH DI LAPTOP (FILE JSON)
        elif os.path.exists("kunci_rahasia.json"):
            
            creds = ServiceAccountCredentials.from_json_keyfile_name("kunci_rahasia.json", scope)
            client = gspread.authorize(creds)
            return client.open_by_key(SPREADSHEET_ID)

        # 3. JIKA KEDUANYA TIDAK ADA
        else:
            st.error("CRITICAL ERROR: Kredensial tidak ditemukan!")
            st.warning("""
            Penyebab:
            1. Jika di Laptop: File 'kunci_rahasia.json' tidak ada di folder ini.
            2. Jika di Cloud: Anda belum setting 'Secrets' di dashboard Streamlit.
            """)
            return None

    except Exception as e:
        st.error(f"GAGAL KONEKSI: {e}")
        return None

@st.cache_data(ttl=300, show_spinner=False)
def get_data(sheet_name):
    sh = connect_to_gsheet()
    if sh:
        try:
            ws = sh.worksheet(sheet_name)
            all_values = ws.get_all_values()
            if len(all_values) > 0:
                headers = all_values[0]
                data = all_values[1:]
                return pd.DataFrame(data, columns=headers)
            else:
                return pd.DataFrame()
        except Exception as e:
            st.error(f"ERROR GSPREAD: {e}")  # Tampilkan Error Asli di Layar
            return pd.DataFrame()
    else:
        st.error("ERROR: Gagal Login ke Google (Kredensial Salah/Tidak Ditemukan)")
        return pd.DataFrame()

# --- FUNGSI KONEKSI & UPLOAD GOOGLE DRIVE ---
@st.cache_resource
def get_drive_service():
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    try:
        if "gcp_service_account" in st.secrets:
            creds_dict = st.secrets["gcp_service_account"]
            creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
        elif os.path.exists("kunci_rahasia.json"):
            creds = ServiceAccountCredentials.from_json_keyfile_name("kunci_rahasia.json", scope)
        else:
            return None
        service = build('drive', 'v3', credentials=creds)
        return service
    except Exception as e:
        return None

def upload_to_drive(file_buffer, filename, mimetype, folder_id):
    service = get_drive_service()
    if not service: return None
    
    file_metadata = {
        'name': filename,
        'parents': [folder_id]
    }
    media = MediaIoBaseUpload(io.BytesIO(file_buffer.getvalue()), mimetype=mimetype, resumable=True)
    
    try:
        # Upload dan minta Google mengembalikan link webViewLink
        file = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink',supportsAllDrives=True).execute()
        return file.get('webViewLink')
    except Exception as e:
        st.error(f"Error API Drive: {e}")
        return None

def save_data(sheet_name, new_data_list):
    """
    Fungsi Reinkarnasi: Membaca seluruh data Sheets, menggabungkan data baru,
    membuang duplikat ID lama, lalu menulis ulang secara bersih dari atas ke bawah.
    100% ANTI NYASAR DAN TAMPILAN SAMA PERSIS DENGAN PORTAL EVALUASI.
    """
    if not new_data_list:
        return True

    sh = connect_to_gsheet()
    if not sh:
        st.error("Tidak bisa konek ke Google Sheets.")
        return False

    try:
        ws = sh.worksheet(sheet_name)
        all_values = ws.get_all_values()

        # 1. Jika Sheet kosong, langsung masukkan data baru beserta headernya
        if not all_values:
            ws.append_rows(new_data_list)
            get_data.clear()
            return True

        headers = all_values[0]
        if sheet_name in ["Price_Data", "Multidrop_Data"] and headers:
            try:
                val_idx = headers.index("validity")
                for row in new_data_list:
                    if len(row) > val_idx:
                        row[val_idx] = str(row[val_idx]).replace(" ", "").strip()
            except ValueError:
                pass
                
        # 2. Buat DataFrame dari data lama di Sheets
        df_old = pd.DataFrame(all_values[1:], columns=headers) if len(all_values) > 1 else pd.DataFrame(columns=headers)
        
        # 3. Buat DataFrame dari data baru yang di-submit vendor
        df_new = pd.DataFrame(new_data_list, columns=headers)

        # 4. Gabungkan data lama dan data baru
        # Kolom pertama headers[0] adalah 'id_transaksi' atau 'id_multidrop'
        key_column = headers[0]
        df_final = pd.concat([df_old, df_new], ignore_index=True)
        
        # Kunci keep='last': Data lama di tengah dihancurkan, data baru masuk otomatis di posisi baris bawah
        df_final = df_final.drop_duplicates(subset=[key_column], keep='last')

        # 5. Bersihkan Sheets dan tulis ulang secara utuh dari baris pertama (A1)
        ws.clear()
        
        # Masukkan kembali header + data final yang sudah bersih tanpa duplikat
        final_matrix = [headers] + df_final.values.tolist()
        ws.update(final_matrix, value_input_option='USER_ENTERED')

        get_data.clear()
        return True

    except Exception as e:
        st.error(f"Gagal memproses: {str(e)}")
        return False

def col_num_to_letter(n):
    """Convert column number to letter. 1=A, 2=B, 26=Z, 27=AA, etc."""
    result = ""
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        result = chr(65 + remainder) + result
    return result

# --- FUNGSI UPDATE STATUS LOCK/UNLOCK (OPTIMIZED) ---
def update_status_locked(ids_to_lock, status_value="Locked"):
    sh = connect_to_gsheet()
    if sh:
        try:
            ws = sh.worksheet("Price_Data")
            vals = ws.get_all_values()
            
            if not vals: return False

            # Cari index kolom
            header = vals[0]
            try:
                id_idx = header.index("id_transaksi")
                status_idx = header.index("status")
            except ValueError:
                return False
            
            # Modifikasi di memori
            is_changed = False
            for i in range(1, len(vals)):
                if vals[i][id_idx] in ids_to_lock:
                    vals[i][status_idx] = status_value
                    is_changed = True
            
            # Upload ulang jika ada perubahan (Batch Update)
            if is_changed:
                ws.update(vals)
            
            get_data.clear()
            return True
        except Exception as e:
            st.error(f"Error update: {e}")
            return False
    return False

# --- ID GENERATORS ---
def generate_next_id(df, id_col, prefix, digits=3):
    if df.empty: return f"{prefix}{str(1).zfill(digits)}"
    existing_ids = [str(x) for x in df[id_col].tolist() if str(x).startswith(prefix)]
    if not existing_ids: return f"{prefix}{str(1).zfill(digits)}"
    max_num = 0
    for x in existing_ids:
        try:
            num_part = x.replace(prefix, "")
            if "-" in num_part: num_part = num_part.split("-")[0]
            val = int(num_part)
            if val > max_num: max_num = val
        except: continue
    return f"{prefix}{str(max_num + 1).zfill(digits)}"

def generate_child_id(df, parent_id, id_col):
    prefix = f"{parent_id}-"
    if df.empty: return f"{prefix}001"
    existing_ids = [str(x) for x in df[id_col].tolist() if str(x).startswith(prefix)]
    if not existing_ids: return f"{prefix}001"
    max_num = 0
    for x in existing_ids:
        try:
            suffix = x.split("-")[-1]
            val = int(suffix)
            if val > max_num: max_num = val
        except: continue
    return f"{prefix}{str(max_num + 1).zfill(3)}"

def clean_numeric(val):
    if pd.isna(val) or str(val).strip() == "": return 0
    try:
        s = str(val).strip()
        # Kalau teksnya berakhiran .0 atau ,0 (desimal gaib Streamlit), potong
        if s.endswith(".0") or s.endswith(",0"):
            s = s[:-2]
            
        s = s.replace("Rp", "").replace(".", "").replace(",", "").replace(" ", "").strip()
        return float(s) if s else 0
    except: 
        return 0

# --- TOMBOL SCROLL TO TOP (VERSI MANUAL ANCHOR) ---
def add_scroll_to_top():
    st.markdown("""
        <style>
            .scroll-to-top {
                position: fixed;
                bottom: 25px;
                right: 25px;
                background-color: #2563EB;
                color: white !important;
                width: 50px;
                height: 50px;
                border-radius: 50%;
                border: 2px solid #1D4ED8;
                text-align: center;
                box-shadow: 2px 2px 10px rgba(0,0,0,0.3);
                cursor: pointer;
                z-index: 99999;
                transition: all 0.3s ease-in-out;
                display: flex;
                align-items: center;
                justify-content: center;
                text-decoration: none; /* Hilangkan garis bawah link */
            }
            .scroll-to-top:hover {
                background-color: #FCA568; 
                border-color: #e38d4a;
                transform: translateY(-5px);
                box-shadow: 2px 5px 15px rgba(252, 165, 104, 0.5);
                color: white !important;
            }
            .scroll-to-top-icon {
                font-size: 24px;
                font-weight: 900;
                line-height: 1;
                margin-bottom: 4px; 
            }
        </style>
        
        <a href="#top-page" class="scroll-to-top" target="_self">
            <span class="scroll-to-top-icon">↑</span>
        </a>
    """, unsafe_allow_html=True)
    
# --- FUNGSI KIRIM EMAIL (UPDATE: ADA INFO ROUND) ---
def send_invitation_email(to_email, vendor_name, load_type, validity, origins, password, round_num="1"):
    # Cek Config
    if "email_config" not in st.secrets:
        st.warning("Konfigurasi email belum disetting di Secrets. Email tidak terkirim.")
        return False

    sender_email = st.secrets["email_config"]["sender_email"]
    sender_password = st.secrets["email_config"]["sender_password"]

    cc_list = ["firli.mandaras@taco.co.id", "budhi.yuono@taco.co.id"]
    cc_string = ", ".join(cc_list)
    
    # Hitung Due Date
    today = datetime.now()
    due_date = today + timedelta(days=7)
    
    months_id = {1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober", 11: "November", 12: "Desember"}
    due_date_str = f"{due_date.day} {months_id[due_date.month]} {due_date.year}"

    # Update Subject 
    subject = f"Undangan Tender {load_type} - {validity} (Tahap Penawaran {round_num}) - TACO Group"
    origins_str = ", ".join(origins)
    
    body = f"""
    <html>
    <body>
        <h3>Dear {vendor_name},</h3>
        <p>Anda telah diundang untuk berpartisipasi dalam Tender Transport <b>TACO Group</b>.</p>
        <p><b>Detail Tender:</b></p>
        <ul>
            <li><b>Periode:</b> {validity}</li>
            <li><b>Tahap Penawaran:</b> {round_num}</li>
            <li><b>Tipe Armada:</b> {load_type}</li>
            <li><b>Area/Origin:</b> {origins_str}</li>
            <li style="color: #d9534f;"><b>Batas Akhir Pengisian: {due_date_str}</b></li>
        </ul>
        <p>Silakan login ke sistem kami untuk memasukkan penawaran harga:</p>
        <p>
            <b>Link App:</b> <a href="https://taco-transport.streamlit.app/">http://bit.ly/TACOtender</a><br>
            <b>Email Login:</b> {to_email}<br>
            <b>Password:</b> {password}<br>
            <b>Tutorial:</b> <a href="https://drive.google.com/file/d/1M5QnSGibg2s9LiQXlNaiCLWxaA7jebJM/view">https://bit.ly/TutorialRFQTACO</a><br>
        </p>
        <p>Terima Kasih,<br>Procurement Team TACO</p>
    </body>
    </html>
    """

    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        msg['Cc'] = cc_string
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html'))

        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        st.error(f"Gagal kirim email: {e}")
        return False

# --- FUNGSI KIRIM EMAIL REMINDER ---
def send_reminder_email(to_email, vendor_name, load_type, validity, round_num, pending_groups, password):
    if "email_config" not in st.secrets: return False
    sender_email = st.secrets["email_config"]["sender_email"]
    sender_password = st.secrets["email_config"]["sender_password"]
    
    cc_list = ["firli.mandaras@taco.co.id", "budhi.yuono@taco.co.id"]
    cc_string = ", ".join(cc_list)
    
    subject = f"REMINDER: Pengisian Penawaran Harga Tender {load_type} - {validity} (Tahap {round_num})"
    
    # PENGAMAN: Ubah semua item ke string 
    pending_groups_str = ", ".join([str(g) for g in pending_groups])
    
    body = f"""
    <html>
    <body>
        <h3 style="color: #d9534f;">⚠️ Reminder Pengisian Penawaran Harga</h3>
        <p>Dear <b>{vendor_name}</b>,</p>
        <p>Melalui email ini, kami ingin mengingatkan bahwa Anda <b>belum menyelesaikan</b> pengisian penawaran harga pada sistem kami untuk detail berikut:</p>
        <ul>
            <li><b>Periode:</b> {validity}</li>
            <li><b>Tahap Penawaran:</b> {round_num}</li>
            <li><b>Tipe Armada:</b> {load_type}</li>
            <li><b style="color: #d9534f;">Origin yang belum diisi: {pending_groups_str}</b></li>
        </ul>
        <p>Mohon segera login dan melengkapi form harga pada area (Origin) yang belum terselesaikan.</p>
        <p>
            <b>Link App:</b> <a href="https://taco-transport.streamlit.app/">http://bit.ly/TACOtender</a><br>
            <b>Email Login:</b> {to_email}<br>
            <b>Password:</b> {password}<br>
            <b>Tutorial:</b> <a href="https://drive.google.com/file/d/1M5QnSGibg2s9LiQXlNaiCLWxaA7jebJM/view">https://bit.ly/TutorialRFQTACO</a><br>
        </p>
        <p>Terima Kasih,<br>Procurement Team TACO</p>
    </body>
    </html>
    """
    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        msg['Cc'] = cc_string 
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html'))
        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
        server.login(sender_email, sender_password)
        server.send_message(msg) 
        server.quit()
        return True
    except Exception as e:
        return False  
        
# --- FUNGSI KIRIM EMAIL REVISI / ALERT ---
def send_rejection_email(to_email, vendor_name, load_type, validity, group_name, reason):
    if "email_config" not in st.secrets: return False
    sender_email = st.secrets["email_config"]["sender_email"]
    sender_password = st.secrets["email_config"]["sender_password"]
    
    # CC ke PIC Internal
    cc_list = ["firli.mandaras@taco.co.id", "budhi.yuono@taco.co.id"]
    cc_string = ", ".join(cc_list)
    
    subject = f"REVISI: Penawaran Harga Tender {load_type} - {validity}"
    
    body = f"""
    <html>
    <body>
        <h3 style="color: #d9534f;">⚠️ Pemberitahuan Revisi Penawaran Harga</h3>
        <p>Dear <b>{vendor_name}</b>,</p>
        <p>Panitia tender TACO Group telah meninjau penawaran harga Anda untuk detail berikut:</p>
        <ul>
            <li><b>Periode:</b> {validity}</li>
            <li><b>Tipe Armada:</b> {load_type}</li>
            <li><b>Group Rute:</b> {group_name}</li>
        </ul>
        <p>Saat ini penawaran Anda <b>membutuhkan revisi</b> dengan catatan sebagai berikut:</p>
        <blockquote style="background-color: #f9f9f9; padding: 10px; border-left: 5px solid #d9534f; margin-left: 0;">
            <i>"{reason}"</i>
        </blockquote>
        <p>Mohon segera login kembali, dan perbaiki sesuai catatan di atas, lalu simpan ulang penawaran Anda.</p>
        <p><b>Link App:</b> <a href="https://taco-transport.streamlit.app/">http://bit.ly/TACOtender</a></p>
        <p>Terima Kasih,<br>Procurement Team TACO</p>
    </body>
    </html>
    """
    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        msg['Cc'] = cc_string 
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html'))
        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
        server.login(sender_email, sender_password)
        server.send_message(msg) 
        server.quit()
        return True
    except Exception as e:
        return False
        
# --- FUNGSI GENERATE WORD SK (UPDATED: 11 KOLOM DENGAN MULTIDROP & BURUH) ---
def create_docx_sk(template_file, nomor_surat, validity, load_type, df_data):
    doc = DocxTemplate(template_file)
    
    # --- 1. SIAPKAN DATA ---
    unique_origins = sorted(df_data['origin'].unique())
    origin_list_str = ", ".join(unique_origins) 

    # --- HELPER 1: SET LEBAR KOLOM ---
    def set_col_widths(table, widths):
        """Mengatur lebar kolom tabel dalam satuan Centimeter (Cm)"""
        table.autofit = False 
        table.allow_autofit = False
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width

    # --- HELPER 2: WARNA BACKGROUND CELL ---
    def set_cell_background(cell, color_hex):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- HELPER 3: FORMAT PARAGRAF ---
    def format_paragraph(paragraph, size, bold=False, align=None):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.line_spacing = 1
        if align is not None: paragraph.alignment = align
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = 'Calibri'

    # --- HELPER 4: REPEAT HEADER ROW ---
    def set_repeat_table_header(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)

    def fmt_val_rp(x):
        try: return f"Rp {int(float(x)):,}".replace(",", ".") if float(x) > 0 else "-"
        except: return "-"

    # --- BAGIAN A: TABEL HARGA ---
    sd = doc.new_subdoc()
    winning_vendors_data = [] 
    
    for org in unique_origins:
        # Judul Origin Area
        p = sd.add_paragraph(f"Origin: {org}")
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0]; run.bold = True; run.font.size = Pt(10)
        
        df_sub = df_data[df_data['origin'] == org].copy()
        df_sub = df_sub.sort_values(by=['kota_asal', 'kota_tujuan', 'unit_type', 'price'])
        winning_vendors_data.append(df_sub)
        
        # Buat Tabel Baru (Melar Jadi 11 Kolom untuk memuat biaya tambahan)
        headers = ['Asal', 'Tujuan', 'Unit', 'Rank', 'Vendor', 'Biaya/Unit', 'Multidrop Dalam', 'Multidrop Luar', 'Biaya Buruh', 'Lead Time', 'TOP']
        table = sd.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER 
        
        # Render Table Header
        hdr_row = table.rows[0]; set_repeat_table_header(hdr_row)
        hdr_cells = hdr_row.cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_background(hdr_cells[i], "ED7D31")
            hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
            format_paragraph(hdr_cells[i].paragraphs[0], size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
        # Render Data Rows
        for _, row in df_sub.iterrows():
            row_obj = table.add_row()         
            row_obj.height = Cm(0.32)         
            row_cells = row_obj.cells
            
            try: harga = f"Rp {int(row['price']):,}".replace(",", ".")
            except: harga = "Rp 0"
              
            lt_raw = str(row['lead_time'])
            lt_fmt = f"{lt_raw} Hari" if lt_raw.isdigit() else "-"
            
            # data multidrop and buruh ke dalam matriks rute Word
            data_map = [
                str(row.get('kota_asal', '-')), 
                str(row['kota_tujuan']), 
                str(row['unit_type']),
                str(row['Ranking']), 
                str(row['vendor_name']), 
                harga,
                fmt_val_rp(row.get('inner_city_price', 0)),
                fmt_val_rp(row.get('outer_city_price', 0)),
                fmt_val_rp(row.get('labor_cost', 0)),
                lt_fmt,
                str(row['top'])
            ]
            
            for idx, val in enumerate(data_map):
                cell = row_cells[idx]; cell.text = val
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # Pengaturan posisi teks angka di kanan, text biasa di kiri, rank di tengah
                if idx in [5, 6, 7, 8]: align = WD_ALIGN_PARAGRAPH.RIGHT
                elif idx in [3, 9, 10]: align = WD_ALIGN_PARAGRAPH.CENTER
                else: align = WD_ALIGN_PARAGRAPH.LEFT
                
                format_paragraph(cell.paragraphs[0], size=6, bold=False, align=align)
        
        # Set Lebar Kolom 11 Kolom secara presisi (Total margin kertas muat rapi)
        col_widths = [Cm(1.49), Cm(2.0), Cm(2.0), Cm(1.0), Cm(2.75), Cm(1.75), Cm(1.56), Cm(1.7), Cm(1.75), Cm(1.0), Cm(1.0)]
        set_col_widths(table, col_widths) 
        sd.add_paragraph("") 

    # --- BAGIAN B: TABEL VENDOR ---
    sd_ven = doc.new_subdoc()
    if winning_vendors_data:
        df_all = pd.concat(winning_vendors_data)
        df_uniq = df_all.drop_duplicates(subset=['vendor_email']).sort_values('vendor_name')
    else: df_uniq = pd.DataFrame()
    
    if not df_uniq.empty:
        v_headers = ['Vendor', 'Alamat', 'Email', 'PIC', 'Telepon']
        v_table = sd_ven.add_table(rows=1, cols=len(v_headers))
        v_table.style = 'Table Grid'
        v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        v_hdr = v_table.rows[0]; set_repeat_table_header(v_hdr)
        vh_cells = v_hdr.cells
        for i, h in enumerate(v_headers):
            vh_cells[i].text = h
            set_cell_background(vh_cells[i], "ED7D31")
            vh_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
            format_paragraph(vh_cells[i].paragraphs[0], size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
        for _, row in df_uniq.iterrows():
            v_cells = v_table.add_row().cells
            v_data = [
                str(row['vendor_name']), 
                str(row.get('address', '-')), 
                str(row['vendor_email']),
                str(row.get('contact_person', '-')), 
                str(row.get('phone', '-'))
            ]
            for idx, val in enumerate(v_data):
                cell = v_cells[idx]; cell.text = val
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                format_paragraph(cell.paragraphs[0], size=7, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)

        v_widths = [Cm(4.0), Cm(5.0), Cm(3.5), Cm(2.5), Cm(2.5)]
        set_col_widths(v_table, v_widths)
   
    bulan_indo = {1:'Januari', 2:'Februari', 3:'Maret', 4:'April', 5:'Mei', 6:'Juni', 7:'Juli', 8:'Agustus', 9:'September', 10:'Oktober', 11:'November', 12:'Desember'}
    today = datetime.now()
    tgl_str = f"{today.day} {bulan_indo[today.month]} {today.year}"
    
    context = {
        'no_surat': nomor_surat,
        'validity': validity,
        'load_type': load_type,
        'tanggal_sk': tgl_str,
        'daftar_origin': origin_list_str,
        'tabel_harga': sd,      
        'tabel_vendor': sd_ven  
    }
    
    doc.render(context)
    output_filename = f"SK_Result_{int(time.time())}.docx"
    doc.save(output_filename)
    return output_filename
    
# --- FUNGSI GENERATE SPH VENDOR ---
def create_docx_sph(template_file, vendor_name, vendor_address, validity, load_type, round_num, df_data):
    doc = DocxTemplate(template_file)
    
    # Format Tanggal
    bulan_indo = {1:'Januari', 2:'Februari', 3:'Maret', 4:'April', 5:'Mei', 6:'Juni', 7:'Juli', 8:'Agustus', 9:'September', 10:'Oktober', 11:'November', 12:'Desember'}
    today = datetime.now()
    tgl_sph = f"{today.day} {bulan_indo[today.month]} {today.year}"

    def fmt_rp(x):
        try: return f"Rp {int(x):,}".replace(",", ".")
        except: return "Rp 0"

    def set_col_widths(table, widths):
        table.autofit = False 
        table.allow_autofit = False
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells): row.cells[idx].width = width

    def format_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=8):
        cell.text = str(text)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        run.font.bold = bold

    def set_repeat_table_header(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)

    sd = doc.new_subdoc()
    
    # --- TABEL 1: HARGA UTAMA ---
    headers = ['No', 'Origin', 'Tujuan', 'Unit', 'Lead Time', 'Harga Penawaran', 'Keterangan']
    table = sd.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shading_elm = parse_xml(r'<w:shd {} w:fill="ED7D31"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        format_cell(hdr_cells[i], h, bold=True, size=8)

    for idx, row in df_data.iterrows():
        row_cells = table.add_row().cells
        
        lt_raw = str(row.get('lead_time', '-'))
        lt_fmt = f"{lt_raw} Hari" if (lt_raw.isdigit() or lt_raw.replace('.','',1).isdigit()) and lt_raw not in ["-","0",""] else "-"
        
        ket = str(row.get('keterangan', '-'))
        if ket.lower() in ["nan", "none", ""]: ket = "-"

        vals = [
            str(idx + 1),
            str(row.get('origin', '-')),
            str(row.get('kota_tujuan', '-')),
            str(row.get('unit_type', '-')),
            lt_fmt,
            fmt_rp(row.get('price', 0)),
            ket
        ]
        for i, v in enumerate(vals):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in [1, 2, 6] else WD_ALIGN_PARAGRAPH.CENTER
            if i == 5: align = WD_ALIGN_PARAGRAPH.RIGHT
            format_cell(row_cells[i], v, align, size=7.5)

    # Lebar Kolom Tabel 1
    col_widths = [Cm(0.8), Cm(2.0), Cm(2.5), Cm(2.0), Cm(1.5), Cm(2.5), Cm(3.0)]
    set_col_widths(table, col_widths)

    sd.add_paragraph("") # Spasi Antar Tabel

    # --- TABEL 2: MULTIDROP & CATATAN ---
    p = sd.add_paragraph("Tabel Biaya Multidrop & Keterangan Tambahan:")
    p.runs[0].bold = True
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(9)
    
    headers_md = ['Origin', 'MD Dalam', 'MD Luar', 'Biaya Buruh', 'Catatan Tambahan Vendor']
    table_md = sd.add_table(rows=1, cols=len(headers_md))
    table_md.style = 'Table Grid'
    table_md.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table_md.rows[0])
    
    hdr_cells_md = table_md.rows[0].cells
    for i, h in enumerate(headers_md):
        shading_elm = parse_xml(r'<w:shd {} w:fill="ED7D31"/>'.format(nsdecls('w')))
        hdr_cells_md[i]._tc.get_or_add_tcPr().append(shading_elm)
        format_cell(hdr_cells_md[i], h, bold=True, size=8)

    # Filter Multidrop agar unik (muncul per Origin saja)
    df_md_uniq = df_data[['origin', 'inner_city_price', 'outer_city_price', 'labor_cost', 'catatan_tambahan']].drop_duplicates(subset=['origin'])
    
    for _, row in df_md_uniq.iterrows():
        row_cells = table_md.add_row().cells
        ctt = str(row.get('catatan_tambahan', '-'))
        if ctt.lower() in ["nan", "none", ""]: ctt = "-"

        vals_md = [
            str(row.get('origin', '-')),
            fmt_rp(row.get('inner_city_price', 0)),
            fmt_rp(row.get('outer_city_price', 0)),
            fmt_rp(row.get('labor_cost', 0)),
            ctt
        ]
        for i, v in enumerate(vals_md):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in [0, 4] else WD_ALIGN_PARAGRAPH.RIGHT
            format_cell(row_cells[i], v, align, size=7.5)
            
    # Lebar Kolom Tabel 2
    col_widths_md = [Cm(2.5), Cm(2.2), Cm(2.2), Cm(2.2), Cm(5.2)]
    set_col_widths(table_md, col_widths_md)

    context = {
        'tanggal': tgl_sph,
        'vendor_name': vendor_name,
        'vendor_address': vendor_address,
        'validity': validity,
        'load_type': load_type,
        'round_num': round_num,
        'tabel_sph': sd
    }
    doc.render(context)
    
    safe_ven = "".join(x for x in vendor_name if x.isalnum())
    fn = f"SPH_{safe_ven}_Tahap{round_num}_{int(time.time())}.docx"
    doc.save(fn)
    return fn
    
# ==============================================================================
# 🎯 FUNGSI UTAMA: LOOP BULK SPK PER VENDOR (PANGGIL DARI TOMBOL STREAMLIT)
# ==============================================================================
def generate_bulk_spk(template_file, nomor_surat, validity, load_type, df_filtered_origin, selected_vendors):
    """
    Fungsi ini dipanggil saat tombol di Streamlit diklik.
    Akan nge-loop setiap vendor yang dipilih dan menghasilkan 1 file Word per vendor.
    """
    import os
    output_folder = "output_spk"
    os.makedirs(output_folder, exist_ok=True)
    
    success_count = 0
    
    # Looping per Vendor yang dipilih di UI Multi-select
    for vendor in selected_vendors:
        # Filter data: Ambil rute yang cuma dimiliki oleh vendor ini
        df_vendor_data = df_filtered_origin[df_filtered_origin['vendor_name'] == vendor].copy()
        
        # Kalau vendor ini gak punya rute di origin yang dipilih, skip biar gak bikin dokumen kosong
        if df_vendor_data.empty:
            continue
            
        # Urutkan data rutenya biar rapi
        df_vendor_data = df_vendor_data.sort_values(by=['kota_asal', 'kota_tujuan', 'unit_type', 'price'])
        
        # Bikin nama file output unik untuk vendor ini
        clean_vendor_name = "".join([c for c in vendor if c.isalpha() or c.isdigit() or c in ' ']).rstrip()
        nama_file_output = os.path.join(output_folder, f"SPK_{clean_vendor_name}_{nomor_surat.replace('/', '_')}.docx")
        
        try:
            # Panggil fungsi generator untuk menggambar tabel ke Word
            create_docx_spk(template_file, nomor_surat, validity, load_type, vendor_name, contact_person, password_last5, origin_combined, alamat_combined, df_data)
            success_count += 1
            st.success(f"✅ Berhasil generate SPK untuk Vendor: **{vendor}**")
        except Exception as e:
            st.error(f"❌ Gagal generate SPK Vendor {vendor}. Error: {str(e)}")
            
    st.info(f"🎉 Selesai! {success_count} dokumen SPK tersimpan di folder `{output_folder}/`")


# ==============================================================================
# 🎯 FUNGSI GENERATOR: DRAW TABEL SPK (9 KOLOM - TANPA RANKING)
# ==============================================================================
def create_docx_spk(template_file, nomor_surat, validity, load_type, vendor_name, contact_person, password_last5, origin_combined, alamat_combined, df_data):
    doc = DocxTemplate(template_file)
    
    # --- HELPER 1: SET LEBAR KOLOM ---
    def set_col_widths(table, widths):
        table.autofit = False 
        table.allow_autofit = False
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width

    # --- HELPER 2: WARNA BACKGROUND CELL ---
    def set_cell_background(cell, color_hex):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # --- HELPER 3: FORMAT PARAGRAF ---
    def format_paragraph(paragraph, size=6, bold=False, align=None):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.line_spacing = 1
        if align is not None: paragraph.alignment = align
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = 'Calibri'

    # --- HELPER 4: REPEAT HEADER ROW ---
    def set_repeat_table_header(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)

    def fmt_val_rp(x):
        try: return f"Rp {int(float(x)):,}".replace(",", ".") if float(x) > 0 else "-"
        except: return "-"

    # --- RENDER TABEL SPK ---
    sd = doc.new_subdoc()
    unique_origins = sorted(df_data['origin'].unique())
    
    for org in unique_origins:
        # Judul Area Origin
        p = sd.add_paragraph(f"Origin: {org}")
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0]; run.bold = True; run.font.size = Pt(10)
        
        df_sub = df_data[df_data['origin'] == org].copy()
        
        # 🎯 HEADERS PAS 9 KOLOM (RANKING REZIMNYA SUDAH TUMBANG/DIHAPUS)
        headers = ['Asal', 'Tujuan', 'Unit', 'Vendor', 'Biaya/Unit', 'Multidrop Dalam', 'Multidrop Luar', 'B.Buruh', 'Lead Time', 'TOP']
        # Sori gais, kolom 'TOP' itu indeks terakhir. Yuk mari di-draw:
        headers = ['Asal', 'Tujuan', 'Unit', 'Vendor', 'Biaya/Unit', 'Multidrop Dalam', 'Multidrop Luar', 'B.Buruh', 'Lead Time', 'TOP']
        
        table = sd.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER 
        
        # Render Table Header
        hdr_row = table.rows[0]; set_repeat_table_header(hdr_row)
        hdr_cells = hdr_row.cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            set_cell_background(hdr_cells[i], "ED7D31")
            hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
            format_paragraph(hdr_cells[i].paragraphs[0], size=7.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            
        # Render Data Rows
        for _, row in df_sub.iterrows():
            row_obj = table.add_row()
            row_obj.height = Cm(0.32)  # 🎯 SAKLEK TINGGI CELL 0.32 CM
            row_cells = row_obj.cells
            
            try: harga = f"Rp {int(row['price']):,}".replace(",", ".")
            except: harga = "Rp 0"
              
            lt_raw = str(row['lead_time'])
            lt_fmt = f"{lt_raw} Hari" if lt_raw.isdigit() else "-"
            
            # 🎯 DATA MAP 9 KOLOM (TANPA RANKING)
            data_map = [
                str(row.get('kota_asal', '-')), 
                str(row['kota_tujuan']), 
                str(row['unit_type']),
                str(row['vendor_name']), 
                harga,
                fmt_val_rp(row.get('inner_city_price', 0)),
                fmt_val_rp(row.get('outer_city_price', 0)),
                fmt_val_rp(row.get('labor_cost', 0)),
                lt_fmt,
                str(row['top'])
            ]
            
            for idx, val in enumerate(data_map):
                cell = row_cells[idx]; cell.text = val
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # 🎯 PENYESUAIAN ALIGNMENT BARU KARENA RANK HILANG
                if idx in [4, 5, 6, 7]: align = WD_ALIGN_PARAGRAPH.RIGHT  # Komponen Harga
                elif idx in [8, 9]: align = WD_ALIGN_PARAGRAPH.CENTER     # Lead Time & TOP
                else: align = WD_ALIGN_PARAGRAPH.LEFT                     # Teks Rute & Vendor
                
                # 🎯 SAKLEK FONT TABEL ISI 6 PT
                format_paragraph(cell.paragraphs[0], size=6, bold=False, align=align)
        
        # 🎯 SAKLEK LEBAR 10 KOLOM BARU SESUAI SELERA REQ LO (TOTAL PAS 18.00 CM)
        col_widths = [
            Cm(1.49),  # Asal
            Cm(2.0),   # Tujuan
            Cm(2.0),   # Unit
            Cm(3.0),   # Vendor
            Cm(1.75),  # Biaya/Unit
            Cm(1.56),  # MD Dalam
            Cm(1.7),   # MD Luar
            Cm(1.75),  # Biaya Buruh
            Cm(1.0),   # Lead Time
            Cm(0.75)   # TOP
        ]
        set_col_widths(table, col_widths) 
        sd.add_paragraph("") 

    bulan_indo = {1:'Januari',2:'Februari',3:'Maret',4:'April',5:'Mei',6:'Juni',
                  7:'Juli',8:'Agustus',9:'September',10:'Oktober',11:'November',12:'Desember'}
    today = datetime.now()
    tgl_spk = f"{today.day} {bulan_indo[today.month]} {today.year}"
    
    context = {
        'no_spk': nomor_surat,
        'validity': validity,
        'load_type': load_type,
        'tanggal_spk': tgl_spk,
        'vendor_name': vendor_name,
        'contact_person': contact_person,
        'password_vendor': password_last5,
        'alamat_gudang': alamat_combined,
        'tabel_harga_vendor': sd
    }
    doc.render(context)
    return doc
    

# ▼ POINTER FIX SAKLEK: TIMPA SELURUH ISI FUNGSI get_target_price DENGAN BLOK INI ▼
def get_target_price(df_all, route_id, unit_type, cur_validity):
    df_safe = df_all.copy()
    if 'price' in df_safe.columns:
        df_safe['price'] = pd.to_numeric(df_safe['price'], errors='coerce').fillna(0)
        df_safe = df_safe[df_safe['price'] > 0]
    else:
        return 0
    
    # 📝 KUNCI SUCI: Normalisasi spasi dan tipe data di memori hitung target price
    df_safe['validity_clean'] = df_safe['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
    df_safe['route_id_clean'] = df_safe['route_id'].astype(str).str.strip()
    df_safe['round_clean'] = pd.to_numeric(df_safe['round'], errors='coerce').fillna(1).astype(int)
    
    clean_cur_val = str(cur_validity).replace(" ", "").lower().strip()
    clean_rid = str(route_id).strip()
    
    # 1. Ambil Harga Terendah Periode SAAT INI (Tahap 1) -> Paksa round == 1 murni angka
    df_curr = df_safe[
        (df_safe['validity_clean'] == clean_cur_val) & 
        (df_safe['route_id_clean'] == clean_rid) & 
        (df_safe['unit_type'].astype(str).str.strip().str.lower() == str(unit_type).strip().lower()) &
        (df_safe['round_clean'] == 1)
    ]
    
    if df_curr.empty: 
        min_curr = 0 
    else:
        min_curr = df_curr['price'].min()

    # 2. Tentukan Periode "SEBELUMNYA" (Histori Masa Lalu)
    target_price = 0
    df_prev = pd.DataFrame()
    
    try:
        parts = str(cur_validity).strip().split(" ") 
        cur_year_str = parts[-1]
        cur_year_int = int(cur_year_str)
        
        is_semester_2 = "juli" in str(cur_validity).lower() or "july" in str(cur_validity).lower()
        
        if is_semester_2:
            df_prev = df_safe[
                (df_safe['validity'].str.contains(cur_year_str, na=False)) & 
                (df_safe['validity'].str.contains("Jan", case=False, na=False)) & 
                (df_safe['route_id_clean'] == clean_rid) & 
                (df_safe['unit_type'].astype(str).str.strip().str.lower() == str(unit_type).strip().str.lower())
            ]
        else:
            prev_year_str = str(cur_year_int - 1)
            df_prev = df_safe[
                (df_safe['validity'].str.contains(prev_year_str, na=False)) & 
                (df_safe['route_id_clean'] == clean_rid) & 
                (df_safe['unit_type'].astype(str).str.strip().str.lower() == str(unit_type).strip().str.lower())
            ]

        # 3. Bandingkan Harga
        if not df_prev.empty:
            min_prev = df_prev['price'].min()
            if min_prev > 0 and min_curr > 0:
                if min_prev < min_curr: target_price = min_prev * 0.95 
                else: target_price = min_curr * 0.85 
            elif min_prev > 0 and min_curr == 0: target_price = min_prev * 0.95
            elif min_prev == 0 and min_curr > 0: target_price = min_curr * 0.85
            else: target_price = 0
        else:
            if min_curr > 0: target_price = min_curr * 0.85
            else: target_price = 0
            
    except:
        if min_curr > 0: target_price = min_curr * 0.85
        else: target_price = 0
        
    return int(target_price)

def generate_session_token(email, password):
    raw = f"{email}:{password}:taco_secret_2025"
    return hashlib.sha256(raw.encode()).hexdigest()[:32]

def try_restore_session(df_users):
    try:
        token_in_url = st.query_params.get("session", "")
        if not token_in_url or df_users.empty:
            return None
        for _, user in df_users.iterrows():
            expected = generate_session_token(user['email'], user['password'])
            if expected == token_in_url:
                return user.to_dict()
    except:
        pass
    return None

def main():
    st.markdown('<div id="top-page"></div>', unsafe_allow_html=True)
    init_style()
    add_scroll_to_top()
    c_logo, _ = st.columns([1, 6])
    with c_logo:
        if os.path.exists("image_2.png"): st.image("image_2.png", width=120)
        else: st.markdown("## **TACO**") 

    if 'user_info' not in st.session_state: st.session_state['user_info'] = None
    if 'vendor_step' not in st.session_state: st.session_state['vendor_step'] = "dashboard" 
    if 'admin_step' not in st.session_state: st.session_state['admin_step'] = "home" 
    # Restore session from URL if page was refreshed
    if st.session_state['user_info'] is None:
        df_users = get_data("Users")
        restored = try_restore_session(df_users)
        if restored:
            st.session_state['user_info'] = restored
    
    for k in ['selected_group_id', 'selected_validity', 'sel_origin', 'sel_load', 'focused_group_id', 'temp_success_msg']:
        if k not in st.session_state: st.session_state[k] = None

    # --- LOGIN ---
    if not st.session_state['user_info']:
        c1, c2, c3 = st.columns([1, 2, 1])
        with c2:
            st.markdown("<br><br>", unsafe_allow_html=True)
            with st.container(border=True):
                st.markdown("### 🔐 TACO Transport RFQ")
                with st.form("login"):
                    email = st.text_input("Email")
                    pw = st.text_input("Password", type="password")
                    if st.form_submit_button("Masuk", type="primary", use_container_width=True):
                        df = get_data("Users")
                        if not df.empty:
                            u = df[(df['email']==email) & (df['password']==pw)]
                            if not u.empty:
                                user_dict = u.iloc[0].to_dict()
                                st.session_state['user_info'] = user_dict
                                token = generate_session_token(email, pw)
                                st.query_params["session"] = token
                                st.rerun()
                            else: st.error("Gagal Login")
                        else: st.error("DB Kosong")
    
    # --- DASHBOARD ---
    else:
        user = st.session_state['user_info']
        role = user['role']

        with st.container():
            c1, c2 = st.columns([8,1])
            with c1: 
                st.markdown(f"### 👋 Halo, **{user.get('vendor_name')}**")
                st.caption(f"Role: {role.upper()}")
            with c2: 
                if st.button("Logout", type="secondary"):
                    st.session_state['user_info'] = None
                    st.session_state['vendor_step'] = "dashboard"
                    st.cache_data.clear()
                    st.query_params.clear()
                    st.rerun()
        
        st.markdown("---")

        if role == 'admin':
            admin_dashboard()
        elif role == 'vendor':
            vendor_dashboard(user['email'])
        else:
            user_dashboard()

# ================= INTERNAL USER DASHBOARD =================
def user_dashboard():
    st.subheader("🔍 Portal Pencarian Tarif")
    
    # --- LOAD DATA ---
    df_p = get_data("Price_Data")
    df_r = get_data("Master_Routes")
    df_g = get_data("Master_Groups")
    df_u = get_data("Users")
    df_prof = get_data("Vendor_Profile")
    df_md = get_data("Multidrop_Data") # Load Multidrop untuk Tab Search

    # --- PREPARE MASTER DATA ---
    df_master = pd.DataFrame()
    if not df_p.empty:
        # Cleaning ID
        df_p['route_id'] = df_p['route_id'].astype(str).str.strip()
        df_r['route_id'] = df_r['route_id'].astype(str).str.strip()
        df_g['group_id'] = df_g['group_id'].astype(str).str.strip()
        if not df_md.empty:
            df_md['vendor_email'] = df_md['vendor_email'].astype(str).str.strip()
            df_md['group_id'] = df_md['group_id'].astype(str).str.strip()
            df_md['validity'] = df_md['validity'].astype(str).str.strip()

        # Merge 1: Price + Routes
        m1 = pd.merge(df_p, df_r, on='route_id', how='left')
        # Merge 2: + Groups
        m2 = pd.merge(m1, df_g, on='group_id', how='left')
        # Merge 3: + User (Vendor Name)
        m3 = pd.merge(m2, df_u[['email', 'vendor_name']], left_on='vendor_email', right_on='email', how='left')
        m3['vendor_name'] = m3['vendor_name'].fillna(m3['vendor_email'])
        
        # Merge 4: + Vendor Profile (TOP, Address, etc)
        if not df_prof.empty:
            df_prof_clean = df_prof.sort_values('updated_at', ascending=False).drop_duplicates('email')
            m4 = pd.merge(m3, df_prof_clean[['email', 'top']], left_on='vendor_email', right_on='email', how='left')
            m4['top'] = m4['top'].fillna("-")
        else:
            m4 = m3
            m4['top'] = "-"

        m4['price'] = pd.to_numeric(m4['price'], errors='coerce').fillna(0)
        
        df_master = m4.copy() 
        
        # --- TAMBAHAN FILTER HARGA 0 ---
        df_master = df_master[df_master['price'] > 0]
    
    # --- TABS ---
    tab1, tab2 = st.tabs(["📊 Summary & Ranking", "🔎 Cari Vendor per Rute"])

    # === TAB 1: SUMMARY RANKING (Mirip Admin) ===
    with tab1:
        if df_master.empty:
            st.info("Data harga belum tersedia.")
        else:
            c1, c2, c3 = st.columns(3)
            # Filter
            avail_val = sorted(df_master['validity'].unique().tolist())
            avail_load = sorted(df_master['load_type'].unique().tolist())
            avail_rounds = sorted(df_master['round'].dropna().unique().tolist())
            
            sel_val = c1.selectbox("Filter Periode", avail_val, key="sum_val")
            sel_load = c2.selectbox("Filter Tipe Muatan", avail_load, key="sum_load")
            sel_round = c3.selectbox("Filter Tahap/Round", avail_rounds, key="sum_round_user", index=len(avail_rounds)-1)
            
            if not df_master.empty:
                df_master_norm = df_master.copy()
                df_master_norm['validity_clean'] = df_master_norm['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                df_master_norm['round_clean'] = pd.to_numeric(df_master_norm['round'], errors='coerce').fillna(1).astype(int)
                    
                    # Bersihkan parameter pemilih selectbox admin dari spasi gaib
                clean_sel_val = str(sel_val).replace(" ", "").lower().strip()
                    
                    # Saring data secara akurat tanpa katarak spasi & tipe data ronde
                df_view = df_master_norm[
                    (df_master_norm['validity_clean'] == clean_sel_val) & 
                    (df_master_norm['load_type'] == sel_load) & 
                    (df_master_norm['round_clean'] == int(sel_round))
                ].copy()
            else:
                df_view = pd.DataFrame()
            
            if not df_view.empty:
                unique_origins = sorted(df_view['origin'].unique())
                for org in unique_origins:
                    with st.expander(f"📍 Origin: {org}", expanded=False):
                        sub_df = df_view[df_view['origin'] == org].copy()
                        # Ranking Logic
                        sub_df = sub_df.sort_values(by=['kota_tujuan', 'unit_type', 'price'])
                        sub_df['Ranking'] = sub_df.groupby(['kota_tujuan', 'unit_type']).cumcount() + 1
                        
                        
                        sub_df['price_fmt'] = sub_df['price'].apply(lambda x: f"Rp {int(x):,}".replace(",", "."))
                        
                        st.dataframe(
                            sub_df[['kota_tujuan', 'unit_type', 'Ranking', 'vendor_name', 'price_fmt', 'lead_time', 'top']],
                            use_container_width=True,
                            hide_index=True,
                            column_config={"kota_tujuan": "Tujuan", "price_fmt": "Harga", "vendor_name": "Vendor", "top": "TOP"}
                        )
            else:
                st.warning("Data tidak ditemukan untuk filter ini.")

    # === TAB 2: SEARCH VENDOR BY ROUTE ===
    with tab2:
        if df_master.empty:
            st.info("Data belum tersedia.")
        else:
            # 1. Filter Utama
            c1, c2, c3 = st.columns(3)
            avail_val_s = sorted(df_master['validity'].unique().tolist())
            s_val = c1.selectbox("1. Pilih Periode", avail_val_s, key="s_val")
            
            avail_load_s = sorted(df_master['load_type'].unique().tolist())
            s_load = c2.selectbox("2. Pilih Muatan", avail_load_s, key="s_load")
            
            filtered_1 = df_master[(df_master['validity'] == s_val) & (df_master['load_type'] == s_load)]
            avail_asal_s = sorted(filtered_1['kota_asal'].dropna().unique().tolist())
            s_org = c3.selectbox("3. Pilih Kota Asal", avail_asal_s, key="s_org")
            
            st.write("")
            df_for_dest = filtered_1[filtered_1['kota_asal'] == s_org].copy()
            avail_dest = sorted(df_for_dest['kota_tujuan'].dropna().unique().tolist())
            
            if avail_dest:
                search_dest = st.selectbox("🔍 Pilih Kota Tujuan", avail_dest, key="s_dest")
                
                if search_dest:
                    df_search = df_for_dest[df_for_dest['kota_tujuan'] == search_dest].copy()
                
                if not df_search.empty:
                    df_search['vendor_email'] = df_search['vendor_email'].astype(str).str.strip().str.lower()
                    df_search['route_id'] = df_search['route_id'].astype(str).str.strip()
                    df_search['validity_clean'] = df_search['validity'].astype(str).str.replace(" ", "").str.replace("-","").str.lower().str.strip()
                    df_search['unit_type'] = df_search['unit_type'].astype(str).str.strip().str.lower()
                    df_search['price'] = pd.to_numeric(df_search['price'], errors='coerce').fillna(0)
    
                    df_search = df_search.sort_values(by=['vendor_email', 'route_id', 'unit_type', 'price'], ascending=True)
                    df_search_clean = df_search.drop_duplicates(subset=['vendor_email', 'route_id', 'unit_type'], keep='first').copy()
                    df_search_clean['group_id_match'] = df_search_clean['route_id'].str[:5].str.upper().str.strip()
    
                    # === PROSES LOOKUP MULTIDROP PERBAIKAN  ===
                    if not df_md.empty:
                        df_md_copy = df_md.copy()
                        df_md_copy['vendor_email_clean'] = df_md_copy['vendor_email'].astype(str).str.strip().str.lower()
                        df_md_copy['validity_clean'] = df_md_copy['validity'].astype(str).str.replace(" ", "").str.replace("-","").str.lower().str.strip()
                        df_md_copy['group_id_clean'] = df_md_copy['group_id'].astype(str).str.upper().str.strip()
                        
                        # Bersihkan nominal dari koma/titik string bawaan Sheets
                        for mc in ['inner_city_price', 'outer_city_price', 'labor_cost']:
                            if mc in df_md_copy.columns:
                                df_md_copy[mc] = df_md_copy[mc].astype(str).str.replace(",", "")
                                df_md_copy[mc] = pd.to_numeric(df_md_copy[mc], errors='coerce').fillna(0)
    
                        df_md_clean = df_md_copy.drop_duplicates(subset=['vendor_email_clean', 'group_id_clean', 'validity_clean'], keep='last')
                        
                        df_result = pd.merge(
                            df_search_clean,
                            df_md_clean[['vendor_email_clean', 'group_id_clean', 'validity_clean', 'inner_city_price', 'outer_city_price', 'labor_cost', 'catatan_tambahan']],
                            left_on=['vendor_email', 'group_id_match', 'validity_clean'],
                            right_on=['vendor_email_clean', 'group_id_clean', 'validity_clean'],
                            how='left'
                        )
                    else:
                        df_result = df_search_clean.copy()
                        df_result['inner_city_price'] = 0
                        df_result['outer_city_price'] = 0
                        df_result['labor_cost'] = 0
                        df_result['catatan_tambahan'] = '-'
    
                    df_result['price'] = pd.to_numeric(df_result['price'], errors='coerce').fillna(0)
                    df_result['inner_city_price'] = pd.to_numeric(df_result['inner_city_price'], errors='coerce').fillna(0)
                    df_result['outer_city_price'] = pd.to_numeric(df_result['outer_city_price'], errors='coerce').fillna(0)
                    df_result['labor_cost'] = pd.to_numeric(df_result['labor_cost'], errors='coerce').fillna(0)
    
                    df_result_display = df_result[df_result['kota_asal'] == s_org].copy()
                    df_result_display = df_result_display.sort_values(by='price', ascending=True)
                    
                    def fmt_rp(x):
                        try: return f"Rp {int(float(x)):,}".replace(",", ".")
                        except: return "Rp 0"
    
                    df_result_display['Harga Unit']      = df_result_display['price'].apply(fmt_rp)
                    df_result_display['Multidrop Dalam'] = df_result_display['inner_city_price'].apply(fmt_rp)
                    df_result_display['Multidrop Luar']  = df_result_display['outer_city_price'].apply(fmt_rp)
                    df_result_display['Biaya Buruh']     = df_result_display['labor_cost'].apply(fmt_rp)
    
                    unique_units = df_result_display['unit_type'].unique()
                    st.success(f"Ditemukan {len(df_result_display)} penawaran untuk tujuan '{search_dest}'.")
    
                    for unit in unique_units:
                        st.markdown(f"##### 🚛 Unit: {unit}")
                        sub_res = df_result_display[df_result_display['unit_type'] == unit].copy().reset_index(drop=True)
                        sub_res['Rank'] = range(1, len(sub_res) + 1)
                        
                        display_cols = ['Rank', 'vendor_name', 'Harga Unit', 'top', 'lead_time', 'Multidrop Dalam', 'Multidrop Luar', 'Biaya Buruh']
                        st.dataframe(
                            sub_res[display_cols],
                            use_container_width=True,
                            hide_index=True,
                            column_config={"vendor_name": "Vendor", "top": "TOP", "lead_time": "Lead Time (Hari)"}
                        )
                        st.markdown("---")
                    else:
                        st.warning(f"Tidak ditemukan rute ke '{search_dest}' dari {s_org}.")
                else:
                    st.info("Silakan ketik nama kota tujuan di atas untuk mulai mencari.")
                
# ================= ADMIN =================
def admin_dashboard():
    step = st.session_state.get('admin_step', 'home')
    
    # --- HALAMAN UTAMA (HOME) ---
    if step == 'home':
        st.markdown("## 🎛️ Admin Portal")
        st.markdown("<br>", unsafe_allow_html=True)
        
        c1, c2 = st.columns(2)
        with c1:
            with st.container(border=True):
                st.markdown("### 📂 Akses & Master Data")
                st.write("Manage Origin, Rute, Unit, User Vendor, dan Akses Pengisian Harga.")
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("Masuk ke Master Data ➡️", type="primary", use_container_width=True):
                    st.session_state['admin_step'] = 'master'
                    st.rerun()
                    
        with c2:
            with st.container(border=True):
                st.markdown("### 📊 Monitoring & Summary")
                st.write("Monitor progres submit vendor, Lock/Unlock data, Ranking Harga, dan Cetak SK/SPK.")
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("Masuk ke Monitoring ➡️", type="primary", use_container_width=True):
                    st.session_state['admin_step'] = 'monitoring'
                    st.rerun()   
                    
    # --- HALAMAN 1: MASTER DATA ---
    elif step == 'master':
        if st.button("⬅️ Kembali ke Menu Utama"):
            st.session_state['admin_step'] = 'home'
            st.rerun() 
        st.markdown("### 📂 Akses & Master Data")
        tabs = st.tabs(["📍Master Groups", "🛣️Master Routes", "🚛Master Units", "👥Users", "🔑Access Rights"])

    
        # --- TAB 1: GROUPS ---
        with tabs[0]:
            st.caption("Buat Group Baru (ID Otomatis)")
            with st.form("add_grp"):
                c1, c2, c3 = st.columns(3)
                lt = c1.selectbox("Load Type", ["FTL", "FCL"])
                org = c2.text_input("Origin (Area)") 
                gn = c3.text_input("Nama Route Group")
            
                if st.form_submit_button("Simpan Group", type="primary"):
                    df_g = get_data("Master_Groups")
                    duplicate = False
                    if not df_g.empty:
                        check = df_g[
                            (df_g['load_type'] == lt) & 
                            (df_g['origin'].str.lower() == org.lower()) & 
                            (df_g['route_group'].str.lower() == gn.lower())
                        ]
                        if not check.empty: duplicate = True
                
                    if duplicate: st.error(f"Gagal: Group '{gn}' dengan Origin '{org}' dan tipe '{lt}' sudah ada!")
                    elif not org or not gn: st.warning("Lengkapi data.")
                    else:
                        new_gid = generate_next_id(df_g, 'group_id', 'R-', 3)
                        save_data("Master_Groups", [[new_gid, lt, gn, org]])
                        st.success(f"Berhasil! ID: {new_gid}")
                        time.sleep(1); st.rerun()
            st.dataframe(get_data("Master_Groups"), use_container_width=True)

        # --- TAB 2: ROUTES ---
        with tabs[1]:
            st.caption("Tambah Rute")
            df_g = get_data("Master_Groups")
        
            c_f1, c_f2 = st.columns(2)
            f_lt = c_f1.selectbox("Filter Load Type", ["All"] + df_g['load_type'].unique().tolist() if not df_g.empty else [])
            avail_origins = []
            if not df_g.empty:
                if f_lt != "All": avail_origins = df_g[df_g['load_type']==f_lt]['origin'].unique().tolist()
                else: avail_origins = df_g['origin'].unique().tolist()
            f_org = c_f2.selectbox("Filter Origin", ["All"] + avail_origins)
        
            grp_opts = {}
            if not df_g.empty:
                filtered_g = df_g.copy()
                if f_lt != "All": filtered_g = filtered_g[filtered_g['load_type'] == f_lt]
                if f_org != "All": filtered_g = filtered_g[filtered_g['origin'] == f_org]
                for _, r in filtered_g.iterrows():
                    label = f"{r['group_id']} | {r['route_group']} ({r['load_type']})"
                    grp_opts[label] = r['group_id']

            with st.form("add_rt"):
                c1, c2, c3 = st.columns([2, 1, 1])
                sel_label = c1.selectbox("Pilih Group", list(grp_opts.keys()) if grp_opts else [])
                ka = c2.text_input("Kota Asal")
                kt = c3.text_input("Kota Tujuan")
                ket = st.text_input("Keterangan (Opsional)")
                
                if st.form_submit_button("Simpan Rute", type="primary"):
                    if sel_label and ka and kt:
                        gid = grp_opts[sel_label]
                        df_r = get_data("Master_Routes")
                        new_rid = generate_child_id(df_r, gid, 'route_id')
                        save_data("Master_Routes", [[new_rid, gid, ka, kt, ket]])
                        st.success(f"Tersimpan! ID: {new_rid}")
                        time.sleep(1); st.rerun()
                    else: st.warning("Data belum lengkap.")
            st.dataframe(get_data("Master_Routes"), use_container_width=True)

        # --- TAB 3: UNITS ---
        with tabs[2]:
            st.caption("Setting Unit per Group (Tanpa Unit ID)")
            all_grp_opts = {}
            if not df_g.empty:
                 for _, r in df_g.iterrows():
                    label = f"{r['group_id']} - {r['route_group']} ({r['origin']})"
                    all_grp_opts[label] = r['group_id']

            with st.form("add_ut"):
                c1, c2 = st.columns(2)
                sel_g = c1.selectbox("Pilih Group", list(all_grp_opts.keys()))
                ut = c2.text_input("Jenis Unit (ex: Tronton)")
            
                if st.form_submit_button("Tambah Unit", type="primary"):
                    if sel_g and ut:
                        gid = all_grp_opts[sel_g]
                        sh = connect_to_gsheet()
                        if sh:
                            ws = sh.worksheet("Master_Units")
                            existing = ws.get_all_values()
                            is_exist = False
                            if len(existing) > 1:
                                for row in existing[1:]:
                                    if len(row) >= 2 and row[0] == gid and row[1].lower() == ut.lower():
                                        is_exist = True; break
                            if is_exist: 
                                st.error(f"Unit '{ut}' sudah ada di Group {gid}.")
                            else:
                                ws.append_rows([[gid, ut]])
                                get_data.clear() 
                                st.success("Unit tersimpan."); time.sleep(0.5); st.rerun()
            st.dataframe(get_data("Master_Units"), use_container_width=True)

        # --- TAB 4: USERS ---
        with tabs[3]:
            with st.form("add_usr"):
                c1, c2, c3 = st.columns(3)
                em = c1.text_input("Email")
                pw = c2.text_input("Pass")
                nm = c3.text_input("PT Name")
                if st.form_submit_button("Add User", type="primary"):
                    save_data("Users", [[em, pw, "vendor", nm]])
                    st.success("Saved")
            
            # Baris ini sekarang sejajar persis dengan 'with st.form'
            st.dataframe(get_data("Users"), use_container_width=True)

# --- TAB 5: ACCESS RIGHTS (UPDATE: FORMAT VENDOR DROPDOWN) ---
        with tabs[4]:
            st.write("Grant Access (Batch per Origin)")
            df_u = get_data("Users"); df_g = get_data("Master_Groups"); df_rights = get_data("Access_Rights")
        
            if not df_u.empty and not df_g.empty:
                # --- HELPER FORMAT DROPDOWN ---
                df_vendors = df_u[df_u['role']=='vendor']
                vendor_emails = df_vendors['email'].unique()
                
                def fmt_vendor(eml):
                    nm = df_vendors[df_vendors['email'] == eml]['vendor_name']
                    if not nm.empty: return f"{nm.iloc[0]} - {eml}"
                    return eml
                # ------------------------------

                c1, c2, c3 = st.columns(3)
                ven = c1.selectbox("Pilih Vendor", vendor_emails, format_func=fmt_vendor) # <-- UPDATE DI SINI
                sel_lt = c2.selectbox("Pilih Load Type", ["FTL", "FCL"])
                sel_round = c3.selectbox("Tahap Penawaran", ["1", "2", "3"]) 
            
                unique_origins = []
                subset_g = df_g[df_g['load_type'] == sel_lt]
                unique_origins = sorted(subset_g['origin'].unique().tolist())
                
                if unique_origins:
                    existing_gids = set()
                    if not df_rights.empty: 
                        if 'round' in df_rights.columns:
                            sub_r = df_rights[(df_rights['vendor_email'] == ven) & (df_rights['round'] == sel_round)]
                        else:
                            sub_r = df_rights[df_rights['vendor_email'] == ven]
                        existing_gids = set(sub_r['group_id'].tolist())

                    with st.form("acc_origin_batch"):
                        cols = st.columns(3)
                        selected_origins = []
                        for idx, org in enumerate(unique_origins):
                            org_groups = df_g[(df_g['origin'] == org) & (df_g['load_type'] == sel_lt)]
                            org_gids = set(org_groups['group_id'].tolist())
                            is_checked = not org_gids.isdisjoint(existing_gids)
                            if cols[idx % 3].checkbox(org, value=is_checked, key=f"chk_{org}"): 
                                selected_origins.append(org)
                        
                        st.divider()
                        c_val1, c_val2 = st.columns([2, 1])
                        
                        opt_validity = ["Januari - Juni", "Juli - Desember"] if sel_lt == "FCL" else ["Januari - Desember"]
                        val_period = c_val1.selectbox("Periode", opt_validity)
                        val_year = c_val2.text_input("Tahun", value=str(datetime.now().year))
                        
                        if st.form_submit_button("Grant Access", type="primary"):
                            val = f"{val_period} {val_year}"
                            if selected_origins:
                                target_groups = df_g[(df_g['load_type'] == sel_lt) & (df_g['origin'].isin(selected_origins))]
                                target_gids = target_groups['group_id'].unique().tolist()
                                
                                sh = connect_to_gsheet()
                                if sh:
                                    ws = sh.worksheet("Access_Rights")
                                    existing_data = ws.get_all_values()
                                    existing_keys = set()
                                    if len(existing_data) > 1:
                                        header = existing_data[0]
                                        has_round_col = 'round' in [h.lower() for h in header]
                                        for row in existing_data[1:]:
                                            if len(row) >= 3:
                                                base_key = f"{row[0]}_{row[1]}_{row[2]}".lower()
                                                if has_round_col and len(row) >= 5: base_key += f"_{row[4]}".lower()
                                                elif not has_round_col: base_key += "_1"
                                                existing_keys.add(base_key)
                                    
                                    new_rows = []
                                    skipped_count = 0
                                    for gid in target_gids:
                                        key_check = f"{ven}_{val}_{gid}_{sel_round}".lower()
                                        if key_check not in existing_keys:
                                            new_rows.append([ven, val, gid, "Active", sel_round])
                                        else: skipped_count += 1
                                    
                                    if new_rows:
                                        ws.append_rows(new_rows)
                                        get_data.clear()
                                        try:
                                            user_row = df_u[df_u['email'] == ven].iloc[0]
                                            with st.spinner("Mengirim email..."):
                                                send_invitation_email(ven, user_row['vendor_name'], sel_lt, val, selected_origins, user_row['password'], sel_round)
                                            st.success(f"Sukses! Akses Tahap {sel_round} diberikan.")
                                        except Exception as e: st.warning(f"Akses OK, Email Gagal: {e}")
                                        time.sleep(1); st.rerun()
                                    else: st.warning(f"Data sudah ada ({skipped_count} skip).")
                            else: st.warning("Pilih minimal 1 origin.")
                    
                    st.markdown("---")
                    # FITUR RESET AKSES
                    with st.expander("🗑️ Area Berbahaya: Reset/Hapus Akses Vendor", expanded=False):
                        c_del1, c_del2 = st.columns(2)
                        del_ven = c_del1.selectbox("Pilih Vendor (Hapus)", vendor_emails, format_func=fmt_vendor, key="del_ven") # <-- UPDATE DI SINI JUGA
                        del_lt = c_del2.selectbox("Pilih Tipe Muatan (Hapus)", ["FTL", "FCL"], key="del_lt")
                        if st.button("⚠️ Hapus Semua Akses Vendor Ini", type="primary"):
                            target_groups_del = df_g[df_g['load_type'] == del_lt]
                            gids_to_remove = set(target_groups_del['group_id'].tolist())
                            sh = connect_to_gsheet()
                            if sh:
                                try:
                                    ws = sh.worksheet("Access_Rights")
                                    all_rows = ws.get_all_values()
                                    if len(all_rows) > 1:
                                        header = all_rows[0]; data_rows = all_rows[1:]
                                        new_data_rows, deleted_count = [], 0
                                        for row in data_rows:
                                            if row[0] == del_ven and row[2] in gids_to_remove: deleted_count += 1
                                            else: new_data_rows.append(row)
                                        if deleted_count > 0:
                                            ws.clear(); ws.append_rows([header] + new_data_rows); get_data.clear()
                                            st.success(f"Dihapus {deleted_count} akses."); time.sleep(1); st.rerun()
                                        else: st.info("Tidak ada data dihapus.")
                                except: st.error("Error Google API")
            
            st.dataframe(get_data("Access_Rights"), use_container_width=True)

    # --- HALAMAN 2: MONITORING & SUMMARY ---
    elif step == 'monitoring':
        if st.button("⬅️ Kembali ke Menu Utama"):
            st.session_state['admin_step'] = 'home'
            st.rerun()
    
        st.markdown("### 📊 Monitoring & Summary")        

        # --- LOAD DATA SEKALI UNTUK SEMUA TAB ANALISA (OPTIMASI) ---

        df_p = get_data("Price_Data")
        df_r = get_data("Master_Routes")
        df_g = get_data("Master_Groups")
        df_u = get_data("Users")
        df_prof = get_data("Vendor_Profile")
        df_md = get_data("Multidrop_Data")
        df_acc = get_data("Access_Rights")
        df_units = get_data("Master_Units")
        df_gudang = get_data("Gudang")
    
        # BIG MERGE MASTER (Untuk Tab 2, 3, 4)
        df_master = pd.DataFrame()
        if not df_p.empty and not df_g.empty:
            df_p['route_id'] = df_p['route_id'].astype(str).str.strip()
            df_r['route_id'] = df_r['route_id'].astype(str).str.strip()
            df_g['group_id'] = df_g['group_id'].astype(str).str.strip()
            m1 = pd.merge(df_p, df_r, on='route_id', how='left')
            m2 = pd.merge(m1, df_g, on='group_id', how='left')
            m3 = pd.merge(m2, df_u[['email', 'vendor_name']], left_on='vendor_email', right_on='email', how='left')
            m3['vendor_name'] = m3['vendor_name'].fillna(m3['vendor_email'])
            if not df_prof.empty:
                df_prof_clean = df_prof.sort_values('updated_at', ascending=False).drop_duplicates('email')
                m4 = pd.merge(m3, df_prof_clean[['email', 'top', 'ppn', 'pph', 'address', 'contact_person', 'phone']], left_on='vendor_email', right_on='email', how='left')
                for c in ['top', 'ppn', 'pph', 'address', 'contact_person', 'phone']:
                    if c in m4.columns: m4[c] = m4[c].fillna("-")
            else:
                m4 = m3
                for c in ['top', 'address', 'contact_person', 'phone']: m4[c] = "-"
            m4['price'] = pd.to_numeric(m4['price'], errors='coerce').fillna(0)
            df_master = m4

        tabs = st.tabs(["⏳ Submit Monitor", "✅ Lock Data", "📊 Summary", "🖨️ Print Dokumen", "📥 SPH Uploads", "Template", "comparison"])
        
# --- TAB 1: SUBMIT MONITOR (UPDATE: STATS & SEARCH BAR) ---
        with tabs[0]:
            if not df_acc.empty and not df_g.empty:
                # Merge Access dengan Group untuk tahu origin, route_group & load type
                acc_merge = pd.merge(df_acc, df_g[['group_id', 'origin', 'route_group', 'load_type']], on='group_id', how='left')
                if 'round' not in acc_merge.columns: acc_merge['round'] = '1'
                
                c1, c2, c3 = st.columns(3)
                sel_sm_lt = c1.selectbox("Filter Tipe Muatan", acc_merge['load_type'].dropna().unique().tolist())
                sel_sm_val = c2.selectbox("Filter Periode", acc_merge['validity'].dropna().unique().tolist())
                sel_sm_rnd = c3.selectbox("Filter Tahap", sorted(acc_merge['round'].dropna().unique().tolist()))
                
                # Filter Data Target (Unik berdasarkan Vendor dan Route Group murni)
                acc_target = acc_merge[
                    (acc_merge['load_type'] == sel_sm_lt) & 
                    (acc_merge['validity'] == sel_sm_val) & 
                    (acc_merge['round'] == sel_sm_rnd)
                ].drop_duplicates(subset=['vendor_email', 'route_group'])
                
                # Ambil Data yg sudah disubmit (dari df_master) dengan normalisasi clean spasi
                sub_master = pd.DataFrame()
                if not df_master.empty:
                    df_m_mon = df_master.copy()
                    df_m_mon['validity_clean'] = df_m_mon['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                    df_m_mon['round_clean'] = pd.to_numeric(df_m_mon['round'], errors='coerce').fillna(1).astype(int)
                    
                    clean_sm_val = str(sel_sm_val).replace(" ", "").lower().strip()
                    
                    sub_master = df_m_mon[
                        (df_m_mon['load_type'] == sel_sm_lt) & 
                        (df_m_mon['validity_clean'] == clean_sm_val) & 
                        (df_m_mon['round_clean'] == int(sel_sm_rnd))
                    ]
                
                if not acc_target.empty:
                    # Tarik database log bypass dari Sheets biar monitor tahu rute mana yang diputihkan
                    df_bp = get_data("Bypass_Monitor")
                    list_bypassed_keys = []
                    if not df_bp.empty:
                        df_bp['bp_key'] = df_bp['vendor_email'].astype(str).str.strip().str.lower() + "_" + \
                                          df_bp['validity'].astype(str).str.replace(" ", "").str.lower().str.strip() + "_" + \
                                          df_bp['round'].astype(str).str.strip()
                        list_bypassed_keys = df_bp['bp_key'].tolist()

                    # --- 1. TAHAP PRE-CALCULATION & PENGUMPULAN DATA ---
                    vendor_data_list = []
                    
                    total_vendors = 0
                    completed_vendors = 0  
                    started_vendors = 0    
                    total_groups_assigned = 0
                    total_groups_filled = 0
                    
                    for vendor in acc_target['vendor_email'].unique():
                        v_name = df_u[df_u['email']==vendor]['vendor_name'].iloc[0] if not df_u[df_u['email']==vendor].empty else vendor
                        
                        v_acc_subset = acc_target[acc_target['vendor_email'] == vendor]
                        assigned_groups = sorted(v_acc_subset['route_group'].dropna().tolist())
                        assigned_origins = v_acc_subset['origin'].dropna().unique().tolist()
                        
                        submitted_groups = []
                        if not sub_master.empty:
                            submitted_groups = sub_master[sub_master['vendor_email'] == vendor]['route_group'].dropna().unique().tolist()
                        
                        # Hitung sisa grup rute asli (sebelum dipotong filter bypass)
                        filled_groups = [grp for grp in assigned_groups if grp in submitted_groups]
                        raw_pending_groups = [grp for grp in assigned_groups if grp not in submitted_groups]
                        
                        # Cek status bypass penawaran selesai manual
                        current_check_key = f"{str(vendor).strip().lower()}_{str(sel_sm_val).replace(' ','').lower().strip()}_{str(sel_sm_rnd).strip()}"
                        
                        is_bypassed = current_check_key in list_bypassed_keys
                        
                        # --- PENENTUAN ANGKA DISPLAY SESUAI KONSEP BARU REKUES LO VAL ---
                        if is_bypassed:
                            # 🎯 JIKA DI-BYPASS: Angka total dan angka terisi dipaksa KEMBAR/SAMA PERSIS!
                            total_g = len(submitted_groups)
                            done_g = len(submitted_groups)
                            pending_groups = [] # Lenyapkan line merah
                        else:
                            # Jika Normal: Ikut hitungan database asli
                            total_g = len(assigned_groups)
                            done_g = len(filled_groups)
                            pending_groups = raw_pending_groups
                        
                        # Hitung Kalkulasi Statistik Lapangan Global
                        total_vendors += 1
                        if len(pending_groups) == 0 and len(assigned_groups) > 0:
                            completed_vendors += 1
                        if len(filled_groups) > 0:
                            started_vendors += 1
                            
                        total_groups_assigned += total_g
                        total_groups_filled += done_g
                        
                        vendor_data_list.append({
                            'email': vendor,
                            'name': v_name,
                            'assigned_groups': assigned_groups,
                            'submitted_groups': submitted_groups,
                            'pending_groups': pending_groups,
                            'raw_pending_groups': raw_pending_groups,
                            'origins': assigned_origins,
                            'total_g': total_g,
                            'done_g': done_g,
                            'is_bypassed': is_bypassed
                        })
                    
                    # Urutkan nama vendor sesuai alfabet biar rapi
                    vendor_data_list.sort(key=lambda x: str(x['name']).strip().lower())
                    
                    # --- 2. TAMPILKAN MATRIKS CARD STATISTIK ---
                    st.divider()
                    col_stat1, col_stat2, col_stat3 = st.columns(3)
                    with col_stat1:
                        st.info(f"🏆 **Selesai** {completed_vendors} / {total_vendors} Vendor")
                    with col_stat2:
                        st.info(f"🏃 **Sudah Mulai Mengisi:** {started_vendors} / {total_vendors} Vendor")
                    with col_stat3:
                        st.info(f"📝 **Grup Rute Sukses Terisi:** {total_groups_filled} / {total_groups_assigned} Grup")
                    
                    # --- 3. SEARCH BAR MONITORING ---
                    search_query = st.text_input("🔍 Cari berdasarkan Nama Vendor atau Origin (Area)...", placeholder="Contoh: Logistik atau Jakarta", key="sm_search_bar_input").strip().lower()
                    st.write("")
                    
                    # --- 4. RENDER COLLAPSE EXPANDER VENDOR ---
                    filtered_vendors = []
                    for v in vendor_data_list:
                        match_name = search_query in str(v['name']).lower()
                        match_origin = any(search_query in str(org).lower() for org in v['origins'])
                        if search_query == "" or match_name or match_origin:
                            filtered_vendors.append(v)
                            
                    if not filtered_vendors:
                        st.warning("Tidak ada nama vendor atau nama area origin yang cocok dengan pencarian.")
                    else:
                        for v in filtered_vendors:
                            v_name = v['name']
                            vendor_email = v['email']
                            assigned_groups = v['assigned_groups']
                            submitted_groups = v['submitted_groups']
                            pending_groups = v['pending_groups']
                            raw_pending_groups = v['raw_pending_groups']
                            total_g = v['total_g']
                            done_g = v['done_g']
                            
                            # 💎 JIKA BYPASS AKTIF: Langsung cetak ikon Ceklis dan paksa angka kembar (misal 3/3 atau 5/5)
                            if v['is_bypassed'] or len(pending_groups) == 0:
                                header_text = f"✅ {v_name} — (Lengkap: {done_g}/{total_g})"
                            else:
                                header_text = f"⚠️ {v_name} — (Selesai: {done_g}/{total_g})"
                                
                            with st.expander(header_text, expanded=False):
                                st.caption(f"📍 Area: {', '.join(v['origins'])}")
                                
                                c_h1, c_h2 = st.columns([4, 2])
                                c_h1.write("**Group Rute**")
                                c_h2.write("**Status Pengisian**")
                                st.divider()
                                
                                for grp in assigned_groups:
                                    # Jika rute ini belum diisi tapi statusnya sudah di-bypass admin, HILANGKAN/SKIP BARISNYA!
                                    if grp in raw_pending_groups and v['is_bypassed']:
                                        continue 
                                        
                                    c_r1, c_r2 = st.columns([4, 2])
                                    c_r1.write(f"{grp}")
                                    if grp in submitted_groups: 
                                        c_r2.markdown('<span class="status-done">✅ Sudah Terisi</span>', unsafe_allow_html=True)
                                    else:
                                        c_r2.markdown('<span class="status-pending">❌ Belum Ada Data</span>', unsafe_allow_html=True)
                                    st.markdown("<hr style='margin: 0.5em 0;'>", unsafe_allow_html=True)
                                
                                # Panel aksi bawah expander
                                if raw_pending_groups and not v['is_bypassed']:
                                    st.write("")
                                    c_btn1, c_btn2 = st.columns(2)
                                    
                                    with c_btn1:
                                        if st.button(f"📨 Kirim Email", key=f"remind_{vendor_email}_{sel_sm_rnd}", type="primary", use_container_width=True):
                                            with st.spinner(f"Mengirim email ke {v_name}..."):
                                                vendor_pw = "Hubungi Admin"
                                                if not df_u[df_u['email']==vendor_email].empty:
                                                    vendor_pw = df_u[df_u['email']==vendor_email].iloc[0]['password']
                                                res = send_reminder_email(vendor_email, v_name, sel_sm_lt, sel_sm_val, sel_sm_rnd, raw_pending_groups, vendor_pw)
                                                if res: st.success("Email terkirim!")
                                                else: st.error("Gagal mengirim email.")
                                    
                                    with c_btn2:
                                        vendor_phone = ""
                                        if not df_prof.empty:
                                            prof_subset = df_prof[df_prof['email'] == vendor_email]
                                            if not prof_subset.empty:
                                                vendor_phone = str(prof_subset.iloc[0].get('phone', '')).strip()
                                        
                                        if vendor_phone and vendor_phone != "-" and vendor_phone.lower() != "nan":
                                            if vendor_phone.startswith("0"):
                                                vendor_phone = "62" + vendor_phone[1:]
                                            pending_str = ", ".join([str(g) for g in raw_pending_groups])
                                            wa_text = f"Halo *{v_name}*,\n\nKami dari TACO Group ingin reminding bahwa Anda *belum menyelesaikan* pengisian harga Tender {sel_sm_lt} ({sel_sm_val}) Tahap {sel_sm_rnd} untuk area:\n\n📌 {pending_str}\n\nMohon segera melengkapi penawaran Anda di sistem.\nLink: https://taco-transport.streamlit.app/\n\n*BATAS PENGISIAN: RABU, 11 MARET 2026*\nJika tidak dilakukan pengisian, kami akan anggap dari {v_name} *TIDAK* akan mengikuti tender rute tersebut dan *TIDAK* dapat menyusul.\n\nTerima Kasih."
                                            wa_text_encoded = urllib.parse.quote(wa_text)
                                            wa_url = f"https://wa.me/{vendor_phone}?text={wa_text_encoded}"
                                            st.markdown(f'<a href="{wa_url}" target="_blank" style="background-color:#25D366; color:white; padding:10px 16px; border-radius:8px; text-decoration:none; font-weight:bold; display:inline-block; text-align:center; width:100%; border: 1px solid #1eaa50; box-shadow: 0 2px 4px rgba(37, 211, 102, 0.2);">💬 Kirim WA</a>', unsafe_allow_html=True)
                                        else:
                                            st.markdown(f'<div style="background-color:#E5E7EB; color:#6B7280; padding:10px 16px; border-radius:8px; text-decoration:none; font-weight:bold; display:inline-block; text-align:center; width:100%; cursor:not-allowed; border: 1px solid #D1D5DB;">❌ No WA Tidak Ada</div>', unsafe_allow_html=True)
                                    
                                    st.write("")
                                    if st.button(f"🔒 Set Selesai Manual", key=f"bypass_{vendor_email}_{sel_sm_rnd}", use_container_width=True, type="secondary"):
                                        with st.spinner("Memproses bypass status vendor..."):
                                            id_bp = f"BP_{vendor_email}_{str(sel_sm_val).replace(' ','')}_{sel_sm_rnd}"
                                            res_bp = save_data("Bypass_Monitor", [[id_bp, vendor_email, sel_sm_val, sel_sm_rnd, "Bypassed", datetime.now().strftime("%Y-%m-%d %H:%M:%S")]])
                                            if res_bp:
                                                st.success(f"Berhasil! Status {v_name} sekarang dianggap selesai manual oleh sistem.")
                                                st.cache_data.clear()
                                                time.sleep(0.5)
                                                st.rerun()
                                else:
                                    st.success("🎉 Pengisian Harga Lengkap!")
                else:
                    st.info("Tidak ada data akses untuk filter ini.")
                    
# --- TAB 2: LOCK DATA ---
        with tabs[1]:
            st.subheader("Lock Data")
            df_price = get_data("Price_Data")
            df_routes = get_data("Master_Routes")
            df_md = get_data("Multidrop_Data")
            df_g = get_data("Master_Groups")
            df_u = get_data("Users") # Tambahan untuk load nama vendor
        
            if df_price.empty:
                st.info("Belum ada data harga masuk.")
            else:
                df_price['route_id'] = df_price['route_id'].astype(str).str.strip()
                df_routes['route_id'] = df_routes['route_id'].astype(str).str.strip()
            
                merged_pr = pd.merge(df_price, df_routes[['route_id', 'group_id', 'kota_asal', 'kota_tujuan']], on='route_id', how='left')
                
                merged_pr['price'] = pd.to_numeric(merged_pr['price'], errors='coerce').fillna(0)
                merged_pr = merged_pr[merged_pr['price'] > 0]
                            
                merged_pr['group_id'] = merged_pr['group_id'].fillna('Unknown')
                
                if not df_g.empty:
                    # Ambil kolom load_type dan origin dari Master Group
                    merged_pr = pd.merge(merged_pr, df_g[['group_id', 'route_group', 'load_type', 'origin']], on='group_id', how='left')
                else: 
                    merged_pr['route_group'] = 'Unknown Group'
                    merged_pr['load_type'] = '-'
                    merged_pr['origin'] = '-'
                    
                merged_pr['route_group'] = merged_pr['route_group'].fillna('Unknown Group')
                merged_pr['load_type'] = merged_pr['load_type'].fillna('-')
                merged_pr['origin'] = merged_pr['origin'].fillna('-')
                merged_pr['kota_asal'] = merged_pr['kota_asal'].fillna('Unknown')
                merged_pr['kota_tujuan'] = merged_pr['kota_tujuan'].fillna('Unknown')
                if 'round' not in merged_pr.columns: merged_pr['round'] = '1'

                # --- MERGE VENDOR NAME UNTUK SEARCH BAR ---
                if not df_u.empty:
                    v_names = df_u[['email', 'vendor_name']]
                    merged_pr = pd.merge(merged_pr, v_names, left_on='vendor_email', right_on='email', how='left')
                    merged_pr['vendor_name'] = merged_pr['vendor_name'].fillna(merged_pr['vendor_email'])
                else:
                    merged_pr['vendor_name'] = merged_pr['vendor_email']

                # ==========================================
                #  FILTER UI LOCK DATA 
                # ==========================================
                c1, c2, c3 = st.columns(3)
                
                avail_round = ["Semua Tahap"] + sorted(merged_pr['round'].unique().tolist())
                sel_ld_round = c1.selectbox("Filter Tahap", avail_round, key="ld_round")
                
                avail_lt = ["Semua Muatan"] + sorted(merged_pr['load_type'].unique().tolist())
                sel_ld_lt = c2.selectbox("Filter Tipe Muatan", avail_lt, key="ld_lt")
                
                avail_val = ["Semua Periode"] + sorted(merged_pr['validity'].unique().tolist())
                sel_ld_val = c3.selectbox("Filter Periode", avail_val, key="ld_val")
                
                search_ld = st.text_input("🔍 Cari Origin atau Nama Vendor...", placeholder="Contoh: Surabaya atau Logistik...", key="ld_search").strip().lower()
                st.write("")
                
                # --- APPLY FILTERS ---
                if sel_ld_round != "Semua Tahap":
                    merged_pr = merged_pr[merged_pr['round'] == sel_ld_round]
                if sel_ld_lt != "Semua Muatan":
                    merged_pr = merged_pr[merged_pr['load_type'] == sel_ld_lt]
                if sel_ld_val != "Semua Periode":
                    merged_pr = merged_pr[merged_pr['validity'] == sel_ld_val]
                    
                if search_ld:
                    match_org = merged_pr['origin'].str.lower().str.contains(search_ld, na=False)
                    match_ven = merged_pr['vendor_name'].str.lower().str.contains(search_ld, na=False)
                    merged_pr = merged_pr[match_org | match_ven]
                # ==========================================
                # ▲▲▲ AKHIR FILTER UI ▲▲▲
                # ==========================================

                if merged_pr.empty:
                    st.info("Tidak ada data yang sesuai dengan filter pencarian.")
                else:
                    # Update key_group agar memuat info tambahan (Tahap dan Nama Vendor)
                    merged_pr['key_group'] = merged_pr['vendor_email'] + " | " + merged_pr['validity'] + " | " + merged_pr['route_group'] + " | " + merged_pr['group_id'] + " | " + merged_pr['vendor_name'] + " | " + merged_pr['round']
                    unique_keys = merged_pr['key_group'].unique()
                    
                    for key in unique_keys:
                        parts = key.split(" | ")
                        vendor, validity, g_name, g_id, v_name_disp, round_disp = parts[0], parts[1], parts[2], parts[3], parts[4], parts[5]
                        subset_pr = merged_pr[merged_pr['key_group'] == key]
                        if subset_pr.empty: continue
                        
                        is_locked = "Locked" in subset_pr['status'].values
                        is_revision = "Need Revision" in subset_pr['status'].values
                        
                        if is_locked: status_icon = "🔒 LOCKED"
                        elif is_revision: status_icon = "⚠️ REVISI"
                        else: status_icon = "🟢 OPEN"
                        
                        l_type = subset_pr.iloc[0]['load_type']
                        
                        # Masukkan ke judul Expander (Sekarang jauh lebih informatif!)
                        with st.expander(f"{status_icon} - Tahap {round_disp} - {l_type} - {v_name_disp} ({validity}) - {g_name}"):
                            st.markdown("**A. Spesifikasi Armada**")
                            if {'unit_type', 'weight_capacity', 'cubic_capacity'}.issubset(subset_pr.columns):
                                df_specs = subset_pr[['unit_type', 'weight_capacity', 'cubic_capacity']].drop_duplicates().reset_index(drop=True)
                                st.dataframe(df_specs, use_container_width=True, hide_index=True)
                    
                            st.markdown("**B. Matriks Harga**")
                            try:
                                subset_pr['price'] = pd.to_numeric(subset_pr['price'], errors='coerce')
                                pivot_df = subset_pr.pivot_table(index=['kota_asal', 'kota_tujuan'], columns='unit_type', values='price', aggfunc='first').reset_index()
                                st.dataframe(pivot_df, use_container_width=True, hide_index=True)
                            except: 
                                st.dataframe(subset_pr, use_container_width=True)

                            st.markdown("**C. Biaya Lain (Multidrop & Buruh)**")
                            if not df_md.empty:
                                df_md['vendor_email'] = df_md['vendor_email'].astype(str).str.strip()
                                df_md['validity'] = df_md['validity'].astype(str).str.strip()
                                df_md['group_id'] = df_md['group_id'].astype(str).str.strip()
                        
                                curr_ven = str(vendor).strip()
                                curr_val = str(validity).strip()
                                curr_gid = str(g_id).strip()

                                sub_md = df_md[
                                    (df_md['vendor_email'] == curr_ven) &
                                    (df_md['validity'] == curr_val) &
                                    (df_md['group_id'] == curr_gid)
                                ]

                                if not sub_md.empty:
                                    cols_to_show = ['inner_city_price', 'outer_city_price']
                                    header_names = ["Multidrop Dalam Kota", " Multidrop Luar Kota"]
                                    
                                    if 'labor_cost' in sub_md.columns:
                                        cols_to_show.append('labor_cost')
                                        header_names.append("Biaya Buruh")
                                    
                                    disp_md = sub_md[cols_to_show].reset_index(drop=True)
                                    disp_md.columns = header_names
                                    st.dataframe(disp_md, use_container_width=True, hide_index=True)

                                    if 'catatan_tambahan' in sub_md.columns:
                                        v_note = str(sub_md.iloc[0]['catatan_tambahan']).strip()
                                        if v_note and v_note.lower() != "nan" and v_note != "None":
                                            st.info(f"📝 **Catatan Tambahan Vendor:**\n{v_note}")
                                else:    
                                    st.info("Data Multidrop belum diinput oleh vendor.")
                            else:
                                st.info("Database Multidrop masih kosong.")
                    
                            # ▼▼▼ BAGIAN D SEKARANG SEJAJAR DENGAN BAGIAN C ▼▼▼
                            st.divider()
                            st.markdown("**D. Alert Revisi ke Vendor**")
                            reason = st.text_input("Catatan Revisi:", key=f"rsn_{key}", placeholder="Contoh: Harga unit Tronton terlalu tinggi, mohon dicek kembali.")
                    
                            c1, c2 = st.columns([1, 3])
                            ids = subset_pr['id_transaksi'].tolist()
                    
                            # Tombol Lock/Unlock
                            if is_locked:
                                if c1.button("🔓 UNLOCK DATA", key=f"ul_{key}"):
                                    update_status_locked(ids, "Open")
                                    st.success("Unlocked!"); time.sleep(0.5); st.rerun()
                            else:
                                if c1.button("🔒 LOCK DATA", key=f"lk_{key}", type="primary"):
                                    update_status_locked(ids, "Locked")
                                    st.success("Locked!"); time.sleep(0.5); st.rerun()
                            
                            # Tombol Kirim Email Alert
                            if c2.button("📨 Kirim Alert Revisi", key=f"rj_{key}"):
                                if not reason:
                                    st.warning("Mohon isi catatan revisi di kotak atas terlebih dahulu.")
                                else:
                                    with st.spinner("Mengirim email alert ke Vendor..."):
                                        v_name = df_u[df_u['email']==vendor]['vendor_name'].iloc[0] if not df_u.empty and not df_u[df_u['email']==vendor].empty else vendor
                                        res = send_rejection_email(vendor, v_name, l_type, validity, g_name, reason)
                                        
                                        if res:
                                            # Ubah status di database menjadi Need Revision
                                            update_status_locked(ids, "Need Revision")
                                            st.success("Berhasil! Email alert terkirim dan status vendor berubah menjadi Revisi.")
                                            time.sleep(1)
                                            st.rerun()
                                        else:
                                            st.error("Gagal mengirim email alert.")
    
        # === TAB 3: SUMMARY & RANKING VENDOR (PATUH TAHAP & AMANKAN TARGET PRICE) ===   
        with tabs[2]:
            st.subheader("📊 Summary & Ranking Vendor")
            
            if df_master.empty:
                st.info("Belum ada data harga masuk.")
            else:
                # --- FILTER UTAMA LAYAR MONITOR ---
                c1, c2, c3, c4 = st.columns(4)
                avail_val = sorted(df_master['validity'].unique().tolist())
                avail_load = sorted(df_master['load_type'].unique().tolist())
                avail_round = sorted(df_master['round'].unique().tolist())
            
                sel_val = c1.selectbox("Filter Periode", avail_val, key="es_val")
                sel_load = c2.selectbox("Filter Tipe Muatan", avail_load, key="es_load")
                # Dropdown filter tahap
                sel_round = c3.selectbox("Filter Tahap", avail_round, key="es_round")
                
                # Saring data kota asal secara dinamis berdasarkan periode and muatan
                df_filter_asal = df_master[(df_master['validity'] == sel_val) & (df_master['load_type'] == sel_load)]
                avail_asal = ["Semua Kota Asal"] + sorted(df_filter_asal['kota_asal'].dropna().unique().tolist())
                sel_asal = c3.selectbox("Filter Kota Asal", avail_asal, key="es_asal") # override kolom c3 untuk kota asal
                
                search_keyword = c4.text_input("🔍 Cari Lokasi", placeholder="Ketik Asal/Tujuan...", key="es_dest").strip().lower()
                
                # --- ENGINE FILTER ---
                df_master_norm = df_master.copy()
                df_master_norm['validity_clean'] = df_master_norm['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                df_master_norm['round_clean'] = pd.to_numeric(df_master_norm['round'], errors='coerce').fillna(1).astype(int)
                
                clean_sel_val = str(sel_val).replace(" ", "").lower().strip()
                
                # Saring data murni hanya yang sesuai ronde and periode di dropdown pilihan
                df_view = df_master_norm[
                    (df_master_norm['validity_clean'] == clean_sel_val) & 
                    (df_master_norm['load_type'] == sel_load) & 
                    (df_master_norm['round_clean'] == int(sel_round))
                ].copy()
                
                df_view = df_view[df_view['price'] > 0]
                
                # --- LANJUTAN FILTER TAMPILAN UI ---
                if sel_asal != "Semua Kota Asal":
                    df_view = df_view[df_view['kota_asal'] == sel_asal]
                    
                if search_keyword:
                    match_org = df_view['origin'].fillna("").str.lower().str.contains(search_keyword)
                    match_asal = df_view['kota_asal'].fillna("").str.lower().str.contains(search_keyword)
                    match_tujuan = df_view['kota_tujuan'].fillna("").str.lower().str.contains(search_keyword)
                    df_view = df_view[match_org | match_asal | match_tujuan]

                # ==========================================================
                # 📥 TAB TABEL EXCEL MASTER SUMMARY (FILTER TAHAP)
                # ==========================================================
                with st.expander("📥 Download Master Summary (Excel)", expanded=False):
                    st.write("Download rekap seluruh rute sesuai filter periode, muatan, and tahap di atas. Rute yang belum diisi vendor akan tetap muncul dengan harga Rp 0.")
                    
                    if not df_r.empty and not df_g.empty and not df_units.empty:
                        # Bersihkan ID
                        df_r_clean = df_r.copy()
                        df_r_clean['route_id'] = df_r_clean['route_id'].astype(str).str.strip()
                        df_r_clean['group_id'] = df_r_clean['group_id'].astype(str).str.strip()
                        df_g_clean = df_g.copy()
                        df_g_clean['group_id'] = df_g_clean['group_id'].astype(str).str.strip()
                        df_u_clean = df_units.copy()
                        df_u_clean['group_id'] = df_u_clean['group_id'].astype(str).str.strip()

                        # Gabungkan Master Kerangka Rute
                        base_df = pd.merge(df_r_clean, df_g_clean, on='group_id', how='left')
                        base_df = pd.merge(base_df, df_u_clean, on='group_id', how='left')
                        
                        # Terapkan Filter UI ke Base Excel
                        base_df = base_df[base_df['load_type'] == sel_load]
                        if sel_asal != "Semua Kota Asal":
                            base_df = base_df[base_df['kota_asal'] == sel_asal]
                        if search_keyword:
                            match_org_b = base_df['origin'].fillna("").str.lower().str.contains(search_keyword)
                            match_asal_b = base_df['kota_asal'].fillna("").str.lower().str.contains(search_keyword)
                            match_tujuan_b = base_df['kota_tujuan'].fillna("").str.lower().str.contains(search_keyword)
                            base_df = base_df[match_org_b | match_asal_b | match_tujuan_b]

                        # Siapkan Harga & Vendor
                        prices_clean = df_p.copy() if not df_p.empty else pd.DataFrame(columns=['route_id', 'unit_type', 'vendor_email', 'price', 'validity', 'round', 'lead_time'])
                        if not prices_clean.empty:
                            prices_clean['route_id'] = prices_clean['route_id'].astype(str).str.strip()
                            prices_clean['price'] = pd.to_numeric(prices_clean['price'], errors='coerce').fillna(0)
                            prices_clean['round_int'] = pd.to_numeric(prices_clean['round'], errors='coerce').fillna(1).astype(int)
                            
                            # 🎯 KUNCI EXCEL: Saring data rekap hanya rute yang harganya valid di ronde pilihan
                            prices_clean = prices_clean[(prices_clean['price'] > 0) & (prices_clean['validity'] == sel_val) & (prices_clean['round_int'] == int(sel_round))]
                            
                            v_names = df_u[df_u['role'] == 'vendor'][['email', 'vendor_name']]
                            prices_clean = pd.merge(prices_clean, v_names, left_on='vendor_email', right_on='email', how='left')
                            prices_clean['vendor_name'] = prices_clean['vendor_name'].fillna(prices_clean['vendor_email'])

                        # Merge Left Join data kerangka rute dengan harga masuk ronde terpilih
                        if not prices_clean.empty:
                            summary_df = pd.merge(
                                base_df, 
                                prices_clean[['route_id', 'unit_type', 'vendor_name', 'price', 'validity', 'round', 'lead_time']], 
                                on=['route_id', 'unit_type'], 
                                how='left'
                            )
                        else:
                            summary_df = base_df.copy()
                            summary_df['vendor_name'] = '-'
                            summary_df['price'] = 0
                            summary_df['validity'] = sel_val
                            summary_df['round'] = sel_round
                            summary_df['lead_time'] = '-'

                        # Hitung Urutan Prioritas Peringkat di Ronde Terpilih
                        summary_df['price_sort'] = summary_df['price'].replace(0, float('inf'))
                        summary_df = summary_df.sort_values(by=['origin', 'kota_asal', 'kota_tujuan', 'unit_type', 'price_sort'])
                        summary_df['Prioritas'] = summary_df.groupby(['origin', 'kota_asal', 'kota_tujuan', 'unit_type']).cumcount() + 1
                        summary_df['Prioritas'] = summary_df.apply(lambda x: x['Prioritas'] if x['price'] > 0 else '-', axis=1)

                        summary_df['price'] = summary_df['price'].fillna(0)
                        summary_df['vendor_name'] = summary_df['vendor_name'].fillna('Belum Ada Penawaran')
                        summary_df['validity'] = summary_df['validity'].fillna(sel_val).replace('nan', sel_val)
                        summary_df['round'] = summary_df['round'].fillna(sel_round).replace('nan', sel_round)
                        summary_df['lead_time'] = summary_df['lead_time'].fillna('-').replace('nan', '-')
                        summary_df['Harga Penawaran'] = summary_df['price'].apply(lambda x: f"Rp {int(x):,}".replace(",", ".") if x > 0 else "Rp 0")
                        
                        cols_to_keep = ['origin', 'kota_asal', 'kota_tujuan', 'route_group', 'load_type', 'unit_type', 'Prioritas', 'vendor_name', 'Harga Penawaran', 'lead_time', 'validity', 'round']
                        for c in cols_to_keep:
                            if c not in summary_df.columns: summary_df[c] = '-'
                        summary_df = summary_df[cols_to_keep]
                        summary_df.columns = ['Origin', 'Kota Asal', 'Kota Tujuan', 'Nama Grup Rute', 'Tipe Muatan', 'Unit', 'Prioritas', 'Nama Vendor', 'Harga Penawaran', 'Lead Time', 'Periode', 'Tahap']
                        
                        # Render file object biner Excel
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            summary_df.to_excel(writer, index=False, sheet_name='Master Summary')
                        excel_data = output.getvalue()
                        
                        safe_val_name = str(sel_val).replace(" - ", "-").replace(" ", "_")
                        st.download_button(
                            label="📊 Download Master Summary (.xlsx)",
                            data=excel_data,
                            file_name=f"Master_Summary_{sel_load}_Tahap{sel_round}_{safe_val_name}_{int(time.time())}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            type="primary",
                            use_container_width=True
                        )
                    else:
                        st.warning("Data Master Rute, Group, atau Unit masih kosong. Tidak bisa mengunduh summary.")
                
                # ==========================================================
                # 🎯 DOWNLOAD TARGET PRICE (EXCEL) 
                # ==========================================================
                with st.expander("🎯 Download Target Price (Excel)", expanded=False):
                    st.write("Download estimasi Target Price untuk fase negosiasi selanjutnya.")
                    
                    if not df_r.empty and not df_g.empty and not df_units.empty:
                        # 1. Siapkan kerangka data (Sama dengan Master Summary)
                        tp_df = pd.merge(df_r_clean, df_g_clean, on='group_id', how='left')
                        tp_df = pd.merge(tp_df, df_u_clean, on='group_id', how='left')
                        
                        # 2. Terapkan Filter yang sedang aktif di UI
                        tp_df = tp_df[tp_df['load_type'] == sel_load]
                        if sel_asal != "Semua Kota Asal":
                            tp_df = tp_df[tp_df['kota_asal'] == sel_asal]
                        if search_keyword:
                            match_org_tp = tp_df['origin'].fillna("").str.lower().str.contains(search_keyword)
                            match_asal_tp = tp_df['kota_asal'].fillna("").str.lower().str.contains(search_keyword)
                            match_tujuan_tp = tp_df['kota_tujuan'].fillna("").str.lower().str.contains(search_keyword)
                            tp_df = tp_df[match_org_tp | match_asal_tp | match_tujuan_tp]

                        # 3. Hitung Target Price Menggunakan Fungsi yang Sudah Ada
                        if st.button("🔄 Generate Data Target Price", key="btn_gen_tp", use_container_width=True):
                            with st.spinner("Menghitung algoritma Target Price untuk seluruh rute..."):
                                tp_results = []
                                prices_for_tp = df_p.copy() if not df_p.empty else pd.DataFrame()
                                
                                for _, row in tp_df.iterrows():
                                    rid = row['route_id']
                                    unit = row['unit_type']
                                    
                                    # Panggil fungsi get_target_price
                                    kalkulasi_tp = get_target_price(prices_for_tp, rid, unit, sel_val)
                                    
                                    tp_results.append({
                                        'Origin': row.get('origin', '-'),
                                        'Kota Asal': row.get('kota_asal', '-'),
                                        'Kota Tujuan': row.get('kota_tujuan', '-'),
                                        'Nama Grup Rute': row.get('route_group', '-'),
                                        'Tipe Muatan': row.get('load_type', '-'),
                                        'Unit': unit,
                                        'Target Price (Angka)': kalkulasi_tp,
                                        'Target Price (Format)': f"Rp {int(kalkulasi_tp):,}".replace(",", ".") if kalkulasi_tp > 0 else "Belum Ada Data"
                                    })
                                
                                df_tp_export = pd.DataFrame(tp_results)
                                
                                # 4. Bikin Excel target price
                                output_tp = io.BytesIO()
                                with pd.ExcelWriter(output_tp, engine='openpyxl') as writer:
                                    df_tp_export.to_excel(writer, index=False, sheet_name='Target Price')
                                excel_tp = output_tp.getvalue()
                                
                                safe_val_name = str(sel_val).replace(" - ", "-").replace(" ", "_")
                                
                                st.download_button(
                                    label="⬇️ Download Target Price (.xlsx)",
                                    data=excel_tp,
                                    file_name=f"Target_Price_{sel_load}_{safe_val_name}_{int(time.time())}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    type="primary",
                                    use_container_width=True
                                )
                    else:
                        st.warning("Data Master masih kosong.")

                # --- 4. TAMPILKAN DISPLAY TABEL SUMMARY DI LAYAR MONITOR ADMIN ---
                if not df_view.empty:
                    unique_origins = sorted(df_view['origin'].unique())
                
                    for org in unique_origins:
                        # Expander tetap per Origin Area agar tetap rapi pengelompokannya
                        with st.expander(f"📍 Origin Area: {org}", expanded=True):
                            sub_df = df_view[df_view['origin'] == org].copy()
                        
                            # Ranking Logic (Diperbarui: Hitung peringkat murni per rute spesifik di ronde pilihan)
                            sub_df = sub_df.sort_values(by=['kota_asal', 'kota_tujuan', 'unit_type', 'price'])
                            sub_df['Ranking'] = sub_df.groupby(['kota_asal', 'kota_tujuan', 'unit_type']).cumcount() + 1
                        

                            
                            sub_df['price_fmt'] = sub_df['price'].apply(lambda x: f"Rp {int(x):,}".replace(",", "."))
                        
                            # Menampilkan 'kota_asal' di dalam tabel layar monitor
                            st.dataframe(
                                sub_df[['kota_asal', 'kota_tujuan', 'unit_type', 'Ranking', 'vendor_name', 'price_fmt', 'lead_time', 'top']],
                                use_container_width=True,
                                column_config={
                                    "kota_asal": "Kota Asal",
                                    "kota_tujuan": "Kota Tujuan",
                                    "unit_type": "Unit",
                                    "price_fmt": "Harga",
                                    "vendor_name": "Vendor",
                                    "top": "TOP"
                                },
                                hide_index=True
                            )
                else:
                    st.warning("Data tidak ditemukan untuk filter kriteria ini.")
                    
# ===================================================================================================

        with tabs[3]:
            st.subheader("🖨️ Print Dokumen")
        
            if df_master.empty:
                st.info("Data belum tersedia.")
            else:
                avail_val = sorted(df_master['validity'].unique().tolist())
                avail_load = sorted(df_master['load_type'].unique().tolist())

                # ==========================================================
                # BAGIAN 1: SURAT KEPUTUSAN (SK) - LIVE RE-MAPPING GITHUB (DUAL ENTITAS)
                # ==========================================================
                with st.container(border=True):
                    st.markdown("### 1. Surat Keputusan (SK)")
                    st.caption("Dokumen rekapitulasi pemenang tender berdasarkan entitas PT otomatis dari server GitHub.")
                
                    # Baris filter dibagi menjadi 4 kolom agar sejajar dan muat dropdown PT nya
                    c1, c2, c3, c4 = st.columns(4)
                    sk_val = c1.selectbox("Periode SK", avail_val, key="sk_val")
                    sk_load = c2.selectbox("Muatan SK", avail_load, key="sk_load")
                    
                    avail_sk_rounds = sorted(df_master['round'].dropna().unique().tolist()) if not df_master.empty else ["1", "2"]
                    sel_sk_round = c3.selectbox("Pilih Tahap SK", avail_sk_rounds, key="sk_round_select", index=len(avail_sk_rounds)-1)
                    
                    # Dropdown pemilih Entitas PT untuk rujukan template SK di GitHub
                    sel_pt_sk = c4.selectbox(
                        "Pilih Entitas Penerbit SK", 
                        ["PT Tangkas Cipta Optimal", "PT Taco Anugrah Corporindo"], 
                        key="sk_pt_entitas_select"
                    )
                
                    # --- PROSES GENERATE DATA FILTER UNTUK SK ---
                    df_master_norm = df_master.copy()
                    df_master_norm['validity_clean'] = df_master_norm['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                    df_master_norm['round_clean'] = pd.to_numeric(df_master_norm['round'], errors='coerce').fillna(1).astype(int)
                    clean_sk_val = str(sk_val).replace(" ", "").lower().strip()
                
                    # Saring basis data kompetisi global seluruh vendor untuk menghitung ranking asli filter tahap
                    df_sk_global = df_master_norm[
                        (df_master_norm['validity_clean'] == clean_sk_val) & 
                        (df_master_norm['load_type'] == sk_load) & 
                        (df_master_norm['round_clean'] == int(sel_sk_round)) &
                        (df_master_norm['price'] > 0)
                    ].copy()
                
                    if not df_sk_global.empty:
                        # Hitung ranking kumulatif global lintas seluruh vendor di tahap terpilih
                        df_sk_global['price_sort_temp'] = df_sk_global['price']
                        df_sk_global = df_sk_global.sort_values(by=['origin', 'kota_asal', 'kota_tujuan', 'unit_type', 'price_sort_temp'])
                        df_sk_global['Ranking'] = df_sk_global.groupby(['origin', 'kota_asal', 'kota_tujuan', 'unit_type']).cumcount() + 1
                        
                        # 🔒 Hitung batas prioritas dinamis berdasarkan angka tertinggi di database (bisa 6,7,8 dst!)
                        max_prio_sk_db = int(df_sk_global['Ranking'].max()) if not df_sk_global.empty else 3
                        prio_options_sk = [i for i in range(1, max_prio_sk_db + 1)]
                        
                        # Sisipkan dropdown sortir prioritas dinamis untuk SK
                        limit_prio_sk = st.selectbox(
                            "🏅 Batasan Urutan Pemenang SK (Sampai Ranking Ke-X)",
                            prio_options_sk,
                            index=len(prio_options_sk)-1, # Default tampilkan semua sampai maksimal peringkat
                            key="sk_prio_limit_select_box"
                        )
                        
                        # Saring data global SK berdasarkan batasan prioritas pilihan
                        df_sk_top_filtered = df_sk_global[df_sk_global['Ranking'] <= limit_prio_sk].copy()
                        
                        avail_org = sorted(df_sk_top_filtered['origin'].unique())
                        sel_orgs = st.multiselect("Pilih Origin Area (SK):", avail_org, default=avail_org, key="sk_orgs")
                
                        if sel_orgs:
                            df_final_sk = df_sk_top_filtered[df_sk_top_filtered['origin'].isin(sel_orgs)].copy()
                            no_sk = st.text_input("Nomor Surat SK:", value="", placeholder="Contoh: 001/SK-PROC/TACO/III/2026", key="no_sk")
                            
                            if st.button("📄 Generate File SK (Otomatis Pisah Per Origin)", type="primary", key="btn_execute_sk_gen"):
                                import requests
                                import zipfile
                                
                                GITHUB_BASE = "https://raw.githubusercontent.com/phibipi/taco-rfq-system/main/templates/"
                                if sel_pt_sk == "PT Tangkas Cipta Optimal":
                                    template_url_sk = GITHUB_BASE + "template_sk_tangkas.docx" 
                                else:
                                    template_url_sk = GITHUB_BASE + "template_sk_tac.docx" 
                                
                                with st.spinner(f"Mengunduh template SK {sel_pt_sk} dan memproses pemecahan origin..."):
                                    try:
                                        response_sk = requests.get(template_url_sk)
                                        if response_sk.status_code != 200:
                                            st.error(f"Gagal mendownload template SK dari GitHub. Status Code: {response_sk.status_code}.")
                                            st.stop()
                                            
                                        # Potong 3 angka urut depan nomor surat resmi
                                        nomor_mentah = str(no_sk).strip()
                                        prefix_angka_str = nomor_mentah[:3]
                                        sisa_nomor_surat = nomor_mentah[3:]
                                        
                                        try:
                                            start_counter = int(prefix_angka_str)
                                            is_numeric_prefix = True
                                        except ValueError:
                                            is_numeric_prefix = False
                                            
                                        zip_buffer = io.BytesIO()
                                        generated_files_count = 0
                                        
                                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                                            # 🎯 LOOPING UTAMA: Pecah berkas terpisah murni per nama origin area
                                            for idx_loop, org_tunggal in enumerate(sorted(sel_orgs)):
                                                df_single_org = df_final_sk[df_final_sk['origin'] == org_tunggal].copy()
                                                
                                                if df_single_org.empty:
                                                    continue
                                                    
                                                # Auto-increment running number nomor surat
                                                if is_numeric_prefix:
                                                    current_num_str = str(start_counter + idx_loop).zfill(3)
                                                    custom_no_sk = f"{current_num_str}{sisa_nomor_surat}"
                                                else:
                                                    custom_no_sk = nomor_mentah
                                                    
                                                tpl_sk_stream = io.BytesIO(response_sk.content)
                                                
                                                # --- ENGINE RE-MAPPING MULTIDROP REAL UNTUK 11 KOLOM SK (FIXED PERMANEN) ---
                                                df_sk_merged = df_single_org.copy()
                                                
                                                if not df_md.empty:
                                                    try:
                                                        df_md_clean = df_md.copy()
                                                        
                                                        # Pastikan semua kolom krusial dikonversi ke string murni sebelum di-strip!
                                                        df_md_clean['vendor_email_clean'] = df_md_clean['vendor_email'].astype(str).str.strip().str.lower()
                                                        df_md_clean['validity_norm'] = df_md_clean['validity'].astype(str).str.replace(" ", "").str.replace("-", "").str.lower().str.strip()
                                                        df_md_clean['group_id_clean'] = df_md_clean['group_id'].astype(str).str.strip().str.upper()
                                                        
                                                        # Normalisasi nominal angka (Buang Rp, titik, koma bawaan Google Sheets)
                                                        for col_num in ['inner_city_price', 'outer_city_price', 'labor_cost']:
                                                            if col_num in df_md_clean.columns:
                                                                df_md_clean[col_num] = df_md_clean[col_num].astype(str).str.replace("Rp", "").str.replace(".", "").str.replace(",", "").str.replace(" ", "").str.strip()
                                                                df_md_clean[col_num] = pd.to_numeric(df_md_clean[col_num], errors='coerce').fillna(0)
                                                        
                                                        # Amankan variabel string filter dari UI Admin (Gunakan str() murni)
                                                        string_sk_val_target = str(sk_val).replace(" ", "").replace("-", "").lower().strip()
                                                        string_sk_round_target = str(sel_sk_round).strip()
                                                        
                                                        # Buat dictionary lookup super cepat
                                                        md_dict_sk = {}
                                                        for _, rmd in df_md_clean.iterrows():
                                                            id_md_raw = str(rmd.get('id_multidrop', '')).strip()
                                                            md_rnd_check = id_md_raw.split("_")[-1] if "_" in id_md_raw else '1'
                                                            
                                                            # Filter validasi kecocokan data multidrop vendor dengan filter pilihan Admin
                                                            if (rmd['validity_norm'] == string_sk_val_target and str(md_rnd_check) == string_sk_round_target):
                                                                k_key = f"{rmd['vendor_email_clean']}_{rmd['group_id_clean']}"
                                                                md_dict_sk[k_key] = {
                                                                    'in': rmd.get('inner_city_price', 0),
                                                                    'out': rmd.get('outer_city_price', 0),
                                                                    'lab': rmd.get('labor_cost', 0)
                                                                }
                                                        
                                                        # Rumus Mapping balik ke matriks data SK
                                                        def lookup_md_to_sk(row_sk):
                                                            v_email = str(row_sk['vendor_email']).strip().lower()
                                                            g_id_raw = row_sk.get('group_id', row_sk['route_id'][:5])
                                                            g_id = str(g_id_raw).strip().upper()
                                                            
                                                            lookup_key = f"{v_email}_{g_id}"
                                                            return md_dict_sk.get(lookup_key, {'in': 0, 'out': 0, 'lab': 0})
                                                        
                                                        # Suntikkan data harga multidrop ke DataFrame final SK
                                                        df_sk_merged['inner_city_price'] = df_sk_merged.apply(lambda x: lookup_md_to_sk(x)['in'], axis=1)
                                                        df_sk_merged['outer_city_price'] = df_sk_merged.apply(lambda x: lookup_md_to_sk(x)['out'], axis=1)
                                                        df_sk_merged['labor_cost'] = df_sk_merged.apply(lambda x: lookup_md_to_sk(x)['lab'], axis=1)
                                                        
                                                    except Exception as ex_sk_md:
                                                        st.error(f"Gagal memproses hitungan biaya tambahan SK: {ex_sk_md}")
                                                        df_sk_merged['inner_city_price'] = 0
                                                        df_sk_merged['outer_city_price'] = 0
                                                        df_sk_merged['labor_cost'] = 0
                                                else:
                                                    df_sk_merged['inner_city_price'] = 0
                                                    df_sk_merged['outer_city_price'] = 0
                                                    df_sk_merged['labor_cost'] = 0
                                                
                                                # Lempar ke fungsi cetak 11 kolom
                                                f_sk_out = create_docx_sk(tpl_sk_stream, custom_no_sk, sk_val, sk_load, df_sk_merged)
                                                
                                                # Inject 3 angka running number ke dalam penamaan file .docx di dalam ZIP
                                                if is_numeric_prefix:
                                                    nomor_urut_file = str(start_counter + idx_loop).zfill(3)
                                                else:
                                                    nomor_urut_file = prefix_angka_str
                                                    
                                                safe_val_sk = str(sk_val).replace(" - ", "-").replace(" ", "_")
                                                safe_pt_name = "Tangkas" if sel_pt_sk == "PT Tangkas Cipta Optimal" else "TAC"
                                                safe_org_name = str(org_tunggal).replace(" ", "")
                                                
                                                filename_word = f"SK_{safe_pt_name}_{safe_org_name}_{sk_load}_{nomor_urut_file}_{safe_val_sk}.docx"
                                                
                                                # Masukkan file tunggal ini ke bungkusan zip, lalu hapus temporary file-nya
                                                zip_file.write(f_sk_out, arcname=filename_word)
                                                os.remove(f_sk_out)
                                                generated_files_count += 1
                                                
                                        if generated_files_count > 0:
                                            safe_val_sk = str(sk_val).replace(" - ", "-").replace(" ", "_")
                                            st.success(f"🎉 Sukses Membelah Data! Berhasil memproduksi {generated_files_count} berkas SK terpisah per origin area.")
                                            st.download_button(
                                                label="⬇️ Download Semua Berkas SK Per Origin (.ZIP)", 
                                                data=zip_buffer.getvalue(), 
                                                file_name=f"SK_{safe_pt_name}_{sk_load}_{safe_val_sk}.zip", 
                                                mime="application/zip",
                                                type="primary",
                                                use_container_width=True
                                            )
                                        else:
                                            st.warning("Tidak ada rute valid yang bisa di-generate.")
                                            
                                    except Exception as e_batch:
                                        st.error(f"Gagal memproses batch split SK: {str(e_batch)}")
                        else:
                            st.warning("Pilih minimal 1 origin area.")
                    else:
                        st.warning("Tidak ditemukan data penawaran harga kompetitor yang aktif untuk kriteria filter ini.")    
                st.write("") # Jarak

               
                # ==========================================================
                # BAGIAN 2: SURAT PERINTAH KERJA (SPK) - MULTI ORIGIN & MULTI VENDOR
                # ==========================================================
                with st.container(border=True):
                    st.markdown("### 2. Surat Perintah Kerja (SPK)")
                    st.caption("Dokumen perintah kerja spesifik per vendor berdasarkan multi-origin, entitas PT, tahap penawaran, dan prioritas global.")
                    
                    # Baris filter utama (Periode, Muatan, Tahap, Entitas PT)
                    c3, c4, c5, c6 = st.columns(4)
                    spk_val = c3.selectbox("Periode SPK", avail_val, key="spk_val")
                    spk_load = c4.selectbox("Muatan SPK", avail_load, key="spk_load")
                    
                    avail_spk_rounds = sorted(df_master['round'].dropna().unique().tolist()) if not df_master.empty else ["1", "2"]
                    sel_spk_round = c5.selectbox("Pilih Tahap SPK", avail_spk_rounds, key="spk_round_select", index=len(avail_spk_rounds)-1)
                    
                    sel_pt_entitas = c6.selectbox(
                        "Pilih Entitas Penerbit SPK", 
                        ["PT Tangkas Cipta Optimal", "PT Taco Anugrah Corporindo"], 
                        key="spk_pt_entitas_select"
                    )

                    # --- PROSES GENERATE LIVE RANKING GLOBAL BERDASARKAN TAHAP ---
                    df_master_norm = df_master.copy()
                    df_master_norm['validity_clean'] = df_master_norm['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                    df_master_norm['round_clean'] = pd.to_numeric(df_master_norm['round'], errors='coerce').fillna(1).astype(int)
                    clean_spk_val = str(spk_val).replace(" ", "").lower().strip()
                    
                    # Filter awal berdasarkan Tahap, Periode, Muatan, dan Harga valid
                    df_spk_global = df_master_norm[
                        (df_master_norm['validity_clean'] == clean_spk_val) & 
                        (df_master_norm['load_type'] == spk_load) & 
                        (df_master_norm['round_clean'] == int(sel_spk_round)) &
                        (df_master_norm['price'] > 0)
                    ].copy()

                    if not df_spk_global.empty:
                        # Hitung ranking kompetisi klasemen antar seluruh vendor (Klasemen Global)
                        df_spk_global['price_sort_temp'] = df_spk_global['price']
                        df_spk_global = df_spk_global.sort_values(by=['origin', 'kota_asal', 'kota_tujuan', 'unit_type', 'price_sort_temp'])
                        df_spk_global['Ranking'] = df_spk_global.groupby(['origin', 'kota_asal', 'kota_tujuan', 'unit_type']).cumcount() + 1
                        
                        st.write("")
                        st.markdown("##### ⚙️ Filter Wilayah & Vendor")
                        
                        # --- 🎯 REVISI UTAMA: BARIS FILTER MULTI-SELECT ORIGIN & VENDOR ---
                        col_spk_o, col_spk_v, col_spk_p = st.columns([1.2, 1.5, 1.3])
                        
                        # 1. Multi-select untuk ORIGIN
                        avail_origins = sorted(df_spk_global['origin'].dropna().unique().tolist())
                        sel_origins = col_spk_o.multiselect("Pilih Rute Origin (Bisa Banyak):", avail_origins, default=avail_origins, key="spk_origin_multiselect")
                        
                        # Filter temporary data berdasarkan Origin yang dipilih untuk memperbarui list Vendor yang tersedia
                        df_spk_filtered_origin = df_spk_global[df_spk_global['origin'].isin(sel_origins)] if sel_origins else df_spk_global
                        
                        # 2. Multi-select untuk VENDOR
                        avail_vens = sorted(df_spk_filtered_origin['vendor_name'].dropna().unique().tolist())
                        sel_vendors = col_spk_v.multiselect("Pilih Vendor Penerima SPK:", avail_vens, key="spk_ven_multiselect")
                        
                        # 3. Dropdown Batasan Prioritas Dinamis
                        max_prio_spk_db = int(df_spk_global['Ranking'].max()) if not df_spk_global.empty else 3
                        prio_options_spk = [i for i in range(1, max_prio_spk_db + 1)]
                        limit_prio_value = col_spk_p.selectbox(
                            "🏅 Batasan Urutan Prioritas Vendor", 
                            prio_options_spk, 
                            index=len(prio_options_spk)-1,
                            key="print_spk_prio_limit_select_box"
                        )
                        
                        no_spk = st.text_input("Nomor Surat SPK (Rujukan Awal):", value="", placeholder="Contoh: 001/SPK/TACO/III/2026", key="no_spk")
                        
                        # --- TOMBOL EKSEKUSI BULK GENERATE SPK (WITH AUTO-INCREMENT NUMBER) ---
                        if st.button("📄 Generate Semua File SPK Vendor", type="primary", key="btn_execute_spk_gen"):
                            if not sel_origins:
                                st.warning("Pilih minimal satu Origin rute dulu dong, gais!")
                            elif not sel_vendors:
                                st.warning("Pilih minimal satu Vendor penerima SPK dulu dong, gais!")
                            elif not no_spk.strip():
                                st.warning("Nomor rujukan awal SPK jangan dikosongin ya!")
                            else:
                                import requests
                                import os
                                import shutil
                                
                                # Siapkan folder output lokal
                                output_folder = "output_spk"
                                os.makedirs(output_folder, exist_ok=True)
                                
                                # --- 🎯 ENGINE AUTO-NUMBER SPK: EKSTRAK 3 ANGKA DEPAN ---
                                nomor_mentah_spk = str(no_spk).strip()
                                prefix_angka_spk = nomor_mentah_spk[:3]
                                sisa_nomor_spk = nomor_mentah_spk[3:]
                                
                                try:
                                    start_counter_spk = int(prefix_angka_spk)
                                    is_numeric_prefix_spk = True
                                except ValueError:
                                    is_numeric_prefix_spk = False
                                
                                GITHUB_BASE = "https://raw.githubusercontent.com/phibipi/taco-rfq-system/main/templates/"
                                template_url = GITHUB_BASE + ("template_spk_tangkas.docx" if sel_pt_entitas == "PT Tangkas Cipta Optimal" else "template_spk_tac.docx")
                                
                                with st.spinner("Mengunduh template dan memproses seluruh SPK dengan auto-number..."):
                                    try:
                                        response = requests.get(template_url)
                                        if response.status_code != 200:
                                            st.error(f"Gagal mendownload template dari GitHub. Status Code: {response.status_code}.")
                                            st.stop()
                                        
                                        success_count = 0
                                        
                                        # 🎯 LOOPING UTAMA: GENERATE 1 FILE PER VENDOR
                                        for idx_loop, v_name in enumerate(sel_vendors):
                                            # Saring data: harus masuk list Origin terpilih, masuk batas prioritas, dan milik Vendor ini
                                            df_final_spk = df_spk_global[
                                                (df_spk_global['origin'].isin(sel_origins)) &
                                                (df_spk_global['Ranking'] <= limit_prio_value) & 
                                                (df_spk_global['vendor_name'] == v_name)
                                            ].copy()
                                            
                                            # Jika kosong (misal vendor ini gak punya rute di origin/prioritas terpilih), skip ke vendor selanjutnya
                                            if df_final_spk.empty:
                                                st.warning(f"⚠️ Vendor **{v_name}** di-skip karena tidak memiliki rute aktif pada Origin terpilih dengan peringkat <= {limit_prio_value}")
                                                continue
                                                
                                            # 🎯 FORMULA AUTO-INCREMENT NOMOR SURAT SPK PER VENDOR REAL TIME:
                                            if is_numeric_prefix_spk:
                                                current_num_spk = str(start_counter_spk + success_count).zfill(3)
                                                custom_no_spk = f"{current_num_spk}{sisa_nomor_spk}"
                                            else:
                                                custom_no_spk = nomor_mentah_spk
                                                
                                            tpl_spk_stream = io.BytesIO(response.content)
                                            
                                            pic = df_final_spk.iloc[0].get('contact_person', 'Pimpinan Perusahaan')
                                            if pd.isna(pic) or pic == "-": pic = "Pimpinan Perusahaan"
                                            
                                            try:
                                                user_row = df_u[df_u['vendor_name'] == v_name]
                                                if not user_row.empty:
                                                    raw_pass = str(user_row.iloc[0]['password'])
                                                    final_pass = raw_pass[-5:] if len(raw_pass) >= 5 else raw_pass
                                                else: final_pass = "XXXXX"
                                            except: final_pass = "XXXXX"

                                            list_origin = sorted(df_final_spk['origin'].unique().tolist())
                                            origin_str_combined = ", ".join(list_origin)
                                            
                                            alamat_list = []
                                            if not df_gudang.empty:
                                                for org in list_origin:
                                                    res_addr = df_gudang[df_gudang['origin'].astype(str).str.lower() == str(org).lower()]
                                                    if not res_addr.empty:
                                                        alamat_found = res_addr.iloc[0]['alamat']
                                                        alamat_list.append(f"{org}: {alamat_found}" if len(list_origin) > 1 else alamat_found)
                                                    else: alamat_list.append(f"{org}: -")
                                            else: alamat_list.append("(Sheet Gudang Kosong)")
                                            alamat_str_combined = "\n".join(alamat_list)

                                            # --- RE-MAPPING LOGIKA MULTIDROP DAN BIAYA BURUH ---
                                            df_spk_merged = df_final_spk.copy()
                                            df_spk_merged['inner_city_price'] = 0
                                            df_spk_merged['outer_city_price'] = 0
                                            df_spk_merged['labor_cost'] = 0
                                            
                                            if not df_md.empty:
                                                try:
                                                    df_md_clean = df_md.copy()
                                                    df_md_clean['vendor_email_clean'] = df_md_clean['vendor_email'].astype(str).str.strip().str.lower()
                                                    df_md_clean['validity_norm'] = df_md_clean['validity'].astype(str).str.replace(" ", "").str.replace("-","").str.lower().str.strip()
                                                    df_md_clean['group_id_clean'] = df_md_clean['group_id'].astype(str).str.upper().str.strip()
                                            
                                                    clean_vendor_email = str(df_final_spk.iloc[0]['vendor_email']).strip().lower()
                                                    clean_validity_spk = str(spk_val).replace(" ", "").replace("-","").lower().strip()
                                            
                                                    md_dict = {}
                                                    for _, rmd in df_md_clean.iterrows():
                                                        id_md_raw = str(rmd.get('id_multidrop', '')).strip()
                                                        md_rnd_check = id_md_raw[-1] if id_md_raw else '1'
                                                            
                                                        if (rmd['vendor_email_clean'] == clean_vendor_email and 
                                                            rmd['validity_norm'] == clean_validity_spk and 
                                                            str(md_rnd_check) == str(sel_spk_round)):
                                                            
                                                            k_gid = rmd['group_id_clean']
                                                            ic_val = str(rmd.get('inner_city_price', '0')).replace(",", "")
                                                            oc_val = str(rmd.get('outer_city_price', '0')).replace(",", "")
                                                            lc_val = str(rmd.get('labor_cost', '0')).replace(",", "")
                                                            
                                                            md_dict[k_gid] = {
                                                                'in': pd.to_numeric(ic_val, errors='coerce') or 0,
                                                                'out': pd.to_numeric(oc_val, errors='coerce') or 0,
                                                                'lab': pd.to_numeric(lc_val, errors='coerce') or 0
                                                            }
                                                    
                                                    def get_md_val(row, kind):
                                                        r_gid = str(row['group_id']).strip().upper()
                                                        res = md_dict.get(r_gid, {'in': 0, 'out': 0, 'lab': 0})
                                                        return res[kind]
                                            
                                                    df_spk_merged['inner_city_price'] = df_spk_merged.apply(lambda x: get_md_val(x, 'in'), axis=1)
                                                    df_spk_merged['outer_city_price'] = df_spk_merged.apply(lambda x: get_md_val(x, 'out'), axis=1)
                                                    df_spk_merged['labor_cost'] = df_spk_merged.apply(lambda x: get_md_val(x, 'lab'), axis=1)
                                                except Exception as ex_md:
                                                    st.error(f"Gagal memproses data biaya tambahan vendor {v_name}: {ex_md}")

                                            # Tentukan nama file output unik per vendor
                                            safe_val = str(spk_val).replace(" - ", "-").replace(" ", "_")
                                            safe_load = str(spk_load).replace(" ", "")
                                            safe_pt_prefix = "TANGKAS" if "Tangkas" in sel_pt_entitas else "TACO"
                                            safe_ven_file = "".join(x for x in v_name if x.isalnum() or x in " -").replace(" ", "_")
                                            
                                            num_file_prefix = current_num_spk if is_numeric_prefix_spk else prefix_angka_spk
                                            custom_filename = f"SPK_{safe_pt_prefix}_{num_file_prefix}_{safe_load}_{safe_val}_{safe_ven_file}.docx"
                                            final_local_path = os.path.join(output_folder, custom_filename)
                                            
                                            # 🎯 BALIKIN PEMANGGILAN ASLI LO BIAR TANGGAL, TABEL, USERNAME GA RUSAK:
                                            doc_obj = create_docx_spk(
                                                tpl_spk_stream, custom_no_spk, spk_val, spk_load,
                                                v_name, pic, final_pass, origin_str_combined,
                                                alamat_str_combined, df_spk_merged
                                            )
                                            
                                            if doc_obj:
                                                doc_obj.save(final_local_path)
                                                success_count += 1
                                                st.write(f"🔹 File sukses dibuat ({custom_no_spk}): `{custom_filename}`")
                                            
                                        # --- JALUR ZIP UTK DOWNLOAD DI LAPTOP ---
                                        if success_count > 0:
                                            import zipfile
                                            
                                            zip_filename = f"ALL_SPK_{safe_load}_{safe_val}.zip"
                                            zip_path = os.path.join(output_folder, zip_filename)
                                            
                                            with zipfile.ZipFile(zip_path, 'w') as zipf:
                                                for root, dirs, files in os.walk(output_folder):
                                                    for file in files:
                                                        if file.endswith('.docx') and not file.startswith('~$'):
                                                            zipf.write(os.path.join(root, file), file)
                                            
                                            st.success(f"🎉 Selesai! Berhasil memproses {success_count} dokumen SPK berurutan di dalam folder `{output_folder}/`!")
                                            
                                            with open(zip_path, "rb") as f_zip:
                                                st.download_button(
                                                    label="📥 DOWNLOAD SEMUA FILE SPK (.ZIP)",
                                                    data=f_zip,
                                                    file_name=zip_filename,
                                                    mime="application/zip",
                                                    type="secondary",
                                                    use_container_width=True
                                                )
                                    except Exception as e: 
                                        st.error(f"Gagal memproses runtutan file Word SPK: {e}")
                    else:
                        st.warning("Tidak ditemukan data penawaran harga kompetitor yang aktif untuk kriteria filter ini.")
        # --- TAB 5: SPH UPLOADS (FITUR BARU) ---
        with tabs[4]:
            st.subheader("📥 Dokumen SPH Vendor (Signed)")
            st.caption("Daftar dokumen Surat Penawaran Harga yang sudah ditandatangani dan di-upload oleh Vendor.")
            
            df_uploads = get_data("SPH_Uploads")
            
            if df_uploads.empty:
                st.info("Belum ada vendor yang mengupload dokumen SPH.")
            else:
                # Tampilkan tabel riwayat upload (terbaru di atas)
                df_uploads = df_uploads.sort_values(by='timestamp', ascending=False)
                
                # Buat tampilan list per dokumen
                for _, row in df_uploads.iterrows():
                    with st.container(border=True):
                        col1, col2, col3 = st.columns([4, 2, 2])
                        col1.markdown(f"**🏢 {row.get('vendor_name', '-')}**")
                        col1.caption(f"Tipe: {row.get('load_type', '-')} | Tahap: {row.get('round', '-')} | Periode: {row.get('validity', '-')}")
                        col2.write(f"🕒 {row.get('timestamp', '-')}")
                        
                        file_url = row.get('filename', '')
                        
                        with col3:
                            if str(file_url).startswith("http"):
                                st.markdown(f'<a href="{file_url}" target="_blank" style="background-color:#2563EB; color:white; padding:8px 12px; border-radius:8px; text-decoration:none; font-weight:bold; display:inline-block; text-align:center; width:100%;">🔗 Buka di Drive</a>', unsafe_allow_html=True)
                            else:
                                st.error("❌ Link Error")               

        # ==================== TAB TEMPLATE (TABS[5]) ====================
        with tabs[5]:
            st.subheader("📄 Template Generator & Pre-populate Horizontal Excel")
            st.caption("Membuat file Excel penawaran harga berjejer ke samping per jenis unit dengan auto-format Currency Rupiah.")

            if not df_g.empty and not df_r.empty and not df_units.empty and not df_u.empty:
                # 🥇 KUNCI SAKLEK: DEKLARASIKAN DF_P_MERGED DI SINI (PALING ATAS) BIAR GAK UNBOUND-LOCAL-ERROR LAGI !
                df_p_merged = pd.DataFrame()
                if not df_master.empty:
                    df_p_merged = df_master.copy()

                # --- 1. FILTER UTAMA TEMPLATE ---
                c1, c2, c3, c4, c5 = st.columns(5)
                
                # Sekarang baris ini aman sentosa karena df_p_merged sudah lahir di atas!
                avail_val_tmpl = sorted(df_p_merged['validity'].dropna().unique().tolist()) if not df_p_merged.empty else ["Januari - Desember 2026"]
                sel_val_tmpl = c1.selectbox("Pilih Periode Laporan", avail_val_tmpl, key="tmpl_val_select")
                
                # 2. Filter Tipe Muatan
                avail_load_tmpl = sorted(df_g['load_type'].unique().tolist())
                sel_load_tmpl = c2.selectbox("Pilih Tipe Muatan", avail_load_tmpl, key="tmpl_load")
                avail_org_tmpl = sorted(df_g[df_g['load_type'] == sel_load_tmpl]['origin'].unique().tolist())
                sel_org_tmpl = c3.multiselect("Pilih Origin Area", avail_org_tmpl, key="tmpl_org")
                
                sel_round_tmpl = c4.selectbox("Pilih Tahap Template", ["Tahap 1", "Tahap 2"], key="tmpl_round_select")
                
                vendor_emails_tmpl = sorted(df_u[df_u['role'] == 'vendor']['email'].unique().tolist())
                def fmt_ven_tmpl(eml):
                    match_name = df_u[df_u['email'] == eml]['vendor_name']
                    if not match_name.empty: return match_name.iloc[0]
                    return eml
                sel_ven_tmpl = c5.selectbox("Pilih Vendor Penerima", vendor_emails_tmpl, format_func=fmt_ven_tmpl, key="tmpl_vendor_select")

                # --- 2. TOMBOL EKSEKUSI GENERATE HORIZONTAL ---
                if st.button("🚀 Generate Template", type="primary", key="btn_run_template_gen"):
                    if not sel_org_tmpl:
                        st.warning("Mohon pilih minimal satu Origin Area terlebih dahulu, Sayang.")
                    else:
                        with st.spinner("Sedang merakit struktur sheet and menyalin histori harga..."):
                            output = io.BytesIO()
                            target_groups = df_g[(df_g['load_type'] == sel_load_tmpl) & (df_g['origin'].isin(sel_org_tmpl))]
                            
                            # Siapkan dataframe rujukan harga yang sudah dibersihkan global
                            df_prices_ref = pd.DataFrame()
                            if not df_p.empty:
                                df_prices_ref = df_p.copy()
                                df_prices_ref['route_id_clean'] = df_prices_ref['route_id'].astype(str).str.strip()
                                df_prices_ref['price'] = pd.to_numeric(df_prices_ref['price'], errors='coerce').fillna(0)
                                df_prices_ref['round_clean_int'] = pd.to_numeric(df_prices_ref['round'], errors='coerce').fillna(1).astype(int)
                                df_prices_ref['vendor_email_clean'] = df_prices_ref['vendor_email'].astype(str).str.strip().str.lower()

                            # Jalankan modul penulisan Excel engine xlsxwriter
                            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                                workbook = writer.book
                                
                                # Pembuatan token format cell Excel Resmi (Mata Uang Rupiah)
                                fmt_header = workbook.add_format({
                                    'bold': True, 
                                    'bg_color': '#FCA568', 
                                    'border': 1, 
                                    'align': 'center', 
                                    'valign': 'vcenter',
                                    'text_wrap': True  # Mengaktifkan bungkusan teks otomatis
                                })
                                fmt_locked = workbook.add_format({'bg_color': '#F3F4F6', 'border': 1, 'align': 'left'})
                                
                                # Format Angka Mata uang: Rp 
                                fmt_currency_locked = workbook.add_format({'num_format': '"Rp "#,##0', 'bg_color': '#F3F4F6', 'border': 1, 'align': 'right'})
                                fmt_currency_input = workbook.add_format({'num_format': '"Rp "#,##0', 'bg_color': '#FFFFFF', 'border': 1, 'align': 'right'})
                                fmt_normal_input = workbook.add_format({'bg_color': '#FFFFFF', 'border': 1, 'align': 'left'})
                                
                                target_email = str(sel_ven_tmpl).lower().strip()

                                for _, g_row in target_groups.iterrows():
                                    gid = g_row['group_id']
                                    g_name = g_row['route_group']
                                    
                                    # Ambil rute unik and jenis unit armada horizontal pada group terkait
                                    routes_sub = df_r[df_r['group_id'] == gid][['route_id', 'kota_asal', 'kota_tujuan']].drop_duplicates().copy()
                                    units_sub = sorted(df_units[df_units['group_id'] == gid]['unit_type'].unique().tolist())
                                    
                                    if not routes_sub.empty and units_sub:
                                        matrix_rows = []
                                        
                                        # Loop per baris Rute
                                        for _, r_row in routes_sub.iterrows():
                                            rid = str(r_row['route_id']).strip()
                                            
                                            row_entry = {
                                                "Route ID": rid,
                                                "Kota Asal": r_row['kota_asal'],
                                                "Kota Tujuan": r_row['kota_tujuan']
                                            }
                                            
                                            # Loop per kolom Jenis Unit (JEJER KE SAMPING)
                                            for unit in units_sub:
                                                if sel_round_tmpl == "Tahap 2":
                                                    # A. Lookup nilai Target Price global rute ini
                                                    tgt_val = get_target_price(df_p_merged, rid, unit, sel_val_tmpl)
                                                    row_entry[f"Target {unit}"] = tgt_val if tgt_val > 0 else 0
                                                    
                                                    # B. Lookup histori Harga Tahap 1 milik vendor ini
                                                    p1_val = 0
                                                    if not df_prices_ref.empty:
                                                        p1_sub = df_prices_ref[
                                                            (df_prices_ref['route_id_clean'] == rid) & 
                                                            (df_prices_ref['unit_type'] == unit) & 
                                                            (df_prices_ref['vendor_email_clean'] == target_email) & 
                                                            (df_prices_ref['round_clean_int'] == 1)
                                                        ]
                                                        if not p1_sub.empty: p1_val = p1_sub['price'].max()
                                                    row_entry[f"Harga Tahap 1 {unit}"] = p1_val
                                                    
                                                    # C. Lookup cicilan Harga Tahap 2 (jika ada)
                                                    p2_val = 0
                                                    if not df_prices_ref.empty:
                                                        p2_sub = df_prices_ref[
                                                            (df_prices_ref['route_id_clean'] == rid) & 
                                                            (df_prices_ref['unit_type'] == unit) & 
                                                            (df_prices_ref['vendor_email_clean'] == target_email) & 
                                                            (df_prices_ref['round_clean_int'] == 2)
                                                        ]
                                                        if not p2_sub.empty: p2_val = p2_sub['price'].max()
                                                    row_entry[f"Harga Tahap 2 {unit}"] = p2_val if p2_val > 0 else ""
                                                else:
                                                    # Jika milih Tahap 1, murni kolom kosong biasa siap ketik
                                                    row_entry[f"Harga Tahap 1 {unit}"] = ""
                                                    
                                            # Tambahkan kolom buntut parameter pendukung di paling ujung kanan
                                            row_entry["Lead Time (Hari)"] = ""
                                            row_entry["Keterangan Vendor"] = ""
                                            matrix_rows.append(row_entry)
                                            
                                        df_sheet_final = pd.DataFrame(matrix_rows)
                                        
                                        # Amankan penamaan sheet dari limit karakater
                                        clean_sheet_name = "".join(x for x in g_name if x.isalnum() or x in " -")[:30]
                                        df_sheet_final.to_excel(writer, sheet_name=clean_sheet_name, index=False)
                                        
                                        worksheet = writer.sheets[clean_sheet_name]
                                        worksheet.set_zoom(55)
                                        worksheet.set_row(0, 30)
                                        
                                        # --- STYLING COLUMN & APPLY CURRENCY RUPIAH SECARA SAKLEK ---
                                        # Atur 4 kolom pembuka wajib dari kiri
                                        worksheet.set_column('A:A', 22, fmt_locked) # Route ID
                                        worksheet.set_column('B:C', 22, fmt_locked) # Asal & Tujuan
                                        
                                        current_col_idx = 3 # Mulai dari kolom index ke-3 (kolom D)
                                        
                                        for unit in units_sub:
                                            if sel_round_tmpl == "Tahap 2":
                                                # Kolom Target Price (Locked & Rupiah format)
                                                worksheet.set_column(current_col_idx, current_col_idx, 18, fmt_currency_locked)
                                                current_col_idx += 1
                                                
                                                # Kolom Harga Tahap 1 (Locked & Rupiah format)
                                                worksheet.set_column(current_col_idx, current_col_idx, 18, fmt_currency_locked)
                                                current_col_idx += 1
                                                
                                                # Kolom Input Harga Tahap 2 (Bisa Diisi & Rupiah format!)
                                                worksheet.set_column(current_col_idx, current_col_idx, 20, fmt_currency_input)
                                                current_col_idx += 1
                                            else:
                                                # Kolom Input Harga Tahap 1 (Bisa Diisi & Rupiah format!)
                                                worksheet.set_column(current_col_idx, current_col_idx, 20, fmt_currency_input)
                                                current_col_idx += 1
                                                
                                        # Atur 2 kolom paling ujung kanan (Lead time and Keterangan)
                                        worksheet.set_column(current_col_idx, current_col_idx, 16, fmt_normal_input) # Lead time
                                        worksheet.set_column(current_col_idx + 1, current_col_idx + 1, 30, fmt_normal_input) # Keterangan
                                        
                                        # Timpa baris paling atas (Header) dengan warna orange
                                        for col_idx, col_name in enumerate(df_sheet_final.columns):
                                            worksheet.write(0, col_idx, col_name, fmt_header)
                                            
                            st.success(f"🎉 Sukses! Berhasil Merakit Template {sel_round_tmpl} untuk {fmt_ven_tmpl(sel_ven_tmpl)}.")
                            
                            safe_ven_fn = str(sel_ven_tmpl).split('@')[0].replace(".","")
                            st.download_button(
                                label=f"⬇️ Ambil File Excel  ({fmt_ven_tmpl(sel_ven_tmpl)})",
                                data=output.getvalue(),
                                file_name=f"Template_{sel_round_tmpl}_{safe_ven_fn}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True
                            )
            else:
                st.error("Database Master tidak lengkap (Master Groups/Routes/Units/Users ada yang kosong).")  

        with tabs[6]: # Tab Perbandingan
            st.subheader("⚖️ Perbandingan Harga Tahap 1 vs Tahap 2")
            st.caption("Lihat penurunan harga per rute dan per vendor untuk negosiasi.")

            if not df_p.empty and not df_r.empty and not df_g.empty:
                # --- 1. DATA PREPARATION (Fixing the Merge issue) ---
                # Pastikan semua ID bersih
                df_p_c = df_p.copy()
                df_r_c = df_r.copy()
                df_g_c = df_g.copy()
                
                df_p_c['route_id'] = df_p_c['route_id'].astype(str).str.strip()
                df_r_c['route_id'] = df_r_c['route_id'].astype(str).str.strip()
                df_r_c['group_id'] = df_r_c['group_id'].astype(str).str.strip()
                df_g_c['group_id'] = df_g_c['group_id'].astype(str).str.strip()

                # Merge Price -> Route -> Group
                df_step1 = pd.merge(df_p_c, df_r_c[['route_id', 'group_id', 'kota_asal', 'kota_tujuan']], on='route_id', how='left')
                df_p_merged = pd.merge(df_step1, df_g_c[['group_id', 'load_type', 'origin']], on='group_id', how='left')
                
                # Paksa kolom kritis jadi numeric
                df_p_merged['price'] = pd.to_numeric(df_p_merged['price'], errors='coerce').fillna(0)
                df_p_merged['round'] = pd.to_numeric(df_p_merged['round'], errors='coerce').fillna(1).astype(int)

                # --- 2. FILTER CONTROLS ---
                c1, c2, c3, c4 = st.columns(4)
                
                avail_val = sorted(df_p_merged['validity'].dropna().unique().tolist())
                sel_val_comp = c1.selectbox("Pilih Periode", avail_val, key="comp_val_final")
                
                avail_lt = sorted(df_p_merged['load_type'].dropna().unique().tolist())
                sel_lt_comp = c2.selectbox("Pilih Tipe Muatan", avail_lt, key="comp_lt_final")
                
                # List Vendor berdasarkan muatan & periode
                vendor_list = sorted(df_p_merged[
                    (df_p_merged['validity'] == sel_val_comp) & 
                    (df_p_merged['load_type'] == sel_lt_comp)
                ]['vendor_email'].dropna().unique().tolist())
                
                if not vendor_list:
                    st.warning("Belum ada data penawaran untuk kriteria ini.")
                else:
                    def fmt_ven_comparison(eml):
                        if not df_u.empty:
                            match_name = df_u[df_u['email'] == eml]['vendor_name']
                            if not match_name.empty:
                                return match_name.iloc[0]
                        return eml

                    sel_ven_comp = c3.selectbox(
                        "Pilih Vendor", 
                        vendor_list, 
                        format_func=fmt_ven_comparison, 
                        key="comp_ven_final"
                    )
                    
                    origin_list = sorted(df_p_merged[
                        (df_p_merged['vendor_email'] == sel_ven_comp) & 
                        (df_p_merged['load_type'] == sel_lt_comp)
                    ]['origin'].dropna().unique().tolist())
                    
                    sel_org_comp = c4.selectbox("Pilih Origin", ["Semua"] + origin_list, key="comp_org_final")

                    # --- 3. PROCESSING COMPARISON ---
                    df_v = df_p_merged[(df_p_merged['vendor_email'] == sel_ven_comp) & (df_p_merged['validity'] == sel_val_comp)]
                    if sel_org_comp != "Semua":
                        df_v = df_v[df_v['origin'] == sel_org_comp]

                    if not df_v.empty:
                        comparison_data = []
                        # Cari rute & unit unik yang pernah diisi vendor ini
                        unique_routes = df_v[['route_id', 'unit_type']].drop_duplicates()
                        
                        for _, r_info in unique_routes.iterrows():
                            rid = r_info['route_id']
                            ut = r_info['unit_type']
                            
                            # Info Lokasi
                            r_row = df_v[df_v['route_id'] == rid].iloc[0]
                            asal_tujuan = f"{r_row['kota_asal']} ➡️ {r_row['kota_tujuan']}"
                            
                            # Harga T1 & T2
                            p1 = df_v[(df_v['route_id'] == rid) & (df_v['unit_type'] == ut) & (df_v['round'] == 1)]['price'].max()
                            p2 = df_v[(df_v['route_id'] == rid) & (df_v['unit_type'] == ut) & (df_v['round'] == 2)]['price'].max()
                            
                            p1 = 0 if pd.isna(p1) else p1
                            p2 = 0 if pd.isna(p2) else p2
                            
                            # Kalkulasi Selisih
                            # Kalkulasi Selisih
                            diff = p1 - p2 if (p1 > 0 and p2 > 0) else 0
                            pct = (diff / p1 * 100) if (p1 > 0 and diff != 0) else 0
                            
                            # ▼ POINTER FIX COMPARISON: HANYA MASUKKAN DATA JIKA TAHAP 1 ATAU TAHAP 2 ADA HARGANYA (BUANG DATA 0 VS 0) ▼
                            if p1 > 0 or p2 > 0:
                                tgt_val = get_target_price(df_p_merged, rid, ut, sel_val_comp)
                                comparison_data.append({
                                    "Origin Area": r_row['origin'],
                                    "Rute": asal_tujuan,
                                    "Unit": ut,
                                    "Harga Tahap 1": p1,
                                    "Harga Tahap 2": p2,
                                    "Target Price": tgt_val,
                                    "Selisih (Rp)": diff,
                                    "Turun (%)": round(pct, 2)
                                })

                        df_final_res = pd.DataFrame(comparison_data)

                        if not df_final_res.empty:
                            # Tampilkan Tabel dengan Warna
                            def color_diff(val):
                                if val > 0: return 'color: green; font-weight: bold'
                                elif val < 0: return 'color: red'
                                return 'color: black'

                            # ▼ POINTER FIX B: DAFTARKAN FORMAT RUPIAH UNTUK KOLOM TARGET PRICE DI TAMPILAN ▼
                            st.dataframe(
                                df_final_res.style.format({
                                    "Harga Tahap 1": "Rp {:,.0f}",
                                    "Harga Tahap 2": "Rp {:,.0f}",
                                    "Target Price": "Rp {:,.0f}",
                                    "Selisih (Rp)": "Rp {:,.0f}",
                                    "Turun (%)": "{:.2f}%"
                                }).map(color_diff, subset=['Selisih (Rp)', 'Turun (%)']),
                                use_container_width=True,
                                hide_index=True
                            )
                            
                            # Summary Statistik
                            total_rute = len(df_final_res)
                            turun_harga = len(df_final_res[df_final_res['Selisih (Rp)'] > 0])
                            st.success(f"📈 Progres: Vendor ini menurunkan harga pada **{turun_harga}** dari **{total_rute}** rute yang ditawarkan.")
                        else:
                            st.info("Tidak ada data harga untuk dibandingkan.")
                    else:
                        st.warning("Data vendor tidak ditemukan.")
            else:
                st.error("Database tidak lengkap (Price/Route/Group missing).")

# ================= VENDOR DASHBOARD (UPDATE: DYNAMIC TABS) =================
def vendor_dashboard(email):
    step = st.session_state['vendor_step']
    
    # --- STEP 1: DASHBOARD / PROFIL ---
    if step == "dashboard":
        # ▼▼▼ CEK STATUS LOCK VENDOR INI ▼▼▼
        df_p_check = get_data("Price_Data")
        has_locked = False
        if not df_p_check.empty and 'status' in df_p_check.columns and 'vendor_email' in df_p_check.columns:
            if "Locked" in df_p_check[df_p_check['vendor_email'] == email]['status'].values:
                has_locked = True
                
        # Bikin Tabs Dinamis
        tab_names = ["🛣️ Pilih Rute & Isi Harga", "📋 Isi Data Perusahaan"]
        if has_locked: 
            tab_names.append("📄 Surat Penawaran Harga")
            
        tabs_ui = st.tabs(tab_names)
        t1 = tabs_ui[0]
        t2 = tabs_ui[1]
        # ▲▲▲ BATAS CEK STATUS ▲▲▲
        
        # Tab 2: Profil
        with t2:
            df_p = get_data("Vendor_Profile")
            curr = {}
            if not df_p.empty:
                m = df_p[df_p['email']==email]
                if not m.empty: curr = m.iloc[-1].to_dict()
            with st.container(border=True):
                with st.form("prof"):
                    c1, c2 = st.columns(2)
                    with c1:
                        ad = st.text_area("Alamat Perusahaan", value=curr.get('address',''))
                        cp = st.text_input("Nama PIC", value=curr.get('contact_person',''))
                    with c2:
                        ph = st.text_input("No. Telepon", value=curr.get('phone',''))
                        top = st.selectbox("Term of Payment", ["7 hari","14 Hari", "30 Hari"])
                    ppn = st.selectbox("PPN", ["11%", "1,1%","0%"])
                    pph = st.selectbox("PPh", ["Include", "Exclude"])
                    if st.form_submit_button("Simpan Data", type="primary"):
                        save_data("Vendor_Profile", [[email, ad, cp, ph, top, ppn, pph, datetime.now().strftime("%Y-%m-%d")]])
                        st.success("Saved")
        
        # Tab 1: List Rute
        with t1:
            df_acc = get_data("Access_Rights")
            df_grps = get_data("Master_Groups")
            df_routes = get_data("Master_Routes")
            df_price = get_data("Price_Data")
            df_gudang = get_data("Gudang")

            if df_acc.empty: 
                st.warning("Belum ada akses.")
                return
            
            my_access = df_acc[df_acc['vendor_email'] == email]
            if my_access.empty: 
                st.info("Anda belum diberikan akses ke project manapun.")
                return

            if 'round' not in my_access.columns: my_access['round'] = '1'
            avail_rounds = sorted(my_access['round'].unique().tolist())
            
            c_filter1, c_filter2 = st.columns(2)
            sel_round = c_filter1.selectbox("Pilih Tahap Penawaran:", avail_rounds)
            my_access = my_access[my_access['round'] == sel_round]

            data_list = []
            for _, acc in my_access.iterrows():
                # PENGAMAN: Gunakan .get()
                gid = acc.get('group_id')
                val = acc.get('validity')
                if not gid or not val: continue
                
                g_info = df_grps[df_grps['group_id'] == gid]
                if not g_info.empty:
                    row = g_info.iloc[0]
                    data_list.append({'validity': val, 'group_id': gid, 'origin': row.get('origin','-'), 'route_group': row.get('route_group','-'), 'load_type': row.get('load_type','-')})
            
            df_disp = pd.DataFrame(data_list)
            if df_disp.empty: 
                st.warning(f"Tidak ada rute di Penawaran Tahap {sel_round}.")
                return

            avail_validity = sorted(df_disp['validity'].unique().tolist())
            sel_val = c_filter2.selectbox("Pilih Periode / Validity:", avail_validity)
            df_view = df_disp[df_disp['validity'] == sel_val]
            
            if df_view.empty: 
                st.info("Tidak ada rute.")
                return
            
            # --- LOGIKA TAB DINAMIS (UPDATE) ---
            # 1. Cek tipe apa saja yang tersedia di data
            avail_types_raw = df_view['load_type'].unique().tolist()
            
            # 2. Definisikan urutan dan nama tab
            type_map = {
                "FTL": "🚛 FTL (Full Truck Load)",
                "FCL": "🚢 FCL (Full Container Load)"
            }
            
            # 3. Filter hanya tipe yang ada datanya (FTL duluan jika ada)
            final_types = [t for t in ["FTL", "FCL"] if t in avail_types_raw]
            
            if not final_types:
                st.warning("Tipe muatan tidak dikenali.")
            else:
                # 4. Buat Tab sesuai isi final_types
                tab_labels = [type_map.get(t, t) for t in final_types]
                created_tabs = st.tabs(tab_labels)
                
                # 5. Loop Render (zip menggabungkan list tipe dan list tab object)
                for t_code, t_ui in zip(final_types, created_tabs):
                    with t_ui:
                        df_sub = df_view[df_view['load_type'] == t_code]
                        # Disini tidak perlu cek empty lagi, karena tab dibuat hanya jika data ada
                        
                        # PENGAMAN: Cegah error origin kosong/tidak terbaca
                        unique_orgs = [o for o in df_sub['origin'].dropna().unique() if str(o).strip() != ""]
                        for org in sorted(unique_orgs, key=lambda x: str(x).strip().lower()):
                            with st.container(border=True):
                                st.markdown(f"#### 📍 {org}")

                                # --- TAMPILKAN ALAMAT GUDANG ---
                                if not df_gudang.empty and 'origin' in df_gudang.columns and 'alamat' in df_gudang.columns:
                                    res_addr = df_gudang[df_gudang['origin'].astype(str).str.lower() == str(org).lower()]
                                    if not res_addr.empty:
                                        st.caption(f"🏢 **Alamat:** {res_addr.iloc[0]['alamat']}")
                                # -------------------------------
                                
                                org_groups = df_sub[df_sub['origin'] == org]
                                c1, c2, c3, c4 = st.columns([3, 4, 2, 2])
                                c1.caption(""); c2.caption("Kota Tujuan"); c3.caption(""); c4.caption("Status Pengisian")
                                st.divider()
                                
                                for _, row in org_groups.iterrows():
                                    gid = row['group_id']
                                    grp_name = row['route_group']
                                    
                                    r_data = df_routes[df_routes['group_id'] == gid] if not df_routes.empty else pd.DataFrame()
                                    status_ui = '<span class="status-pending">❌ Belum Ada Data</span>'
                                    is_locked_btn = False
                                    
                                    
                                    if not df_price.empty and not r_data.empty:
                                        # Buat dataframe normalisasi sementara khusus buat ngecek status doang
                                        df_p_status = df_price.copy()
                                        df_p_status['vendor_email_clean'] = df_p_status['vendor_email'].astype(str).str.strip().str.lower()
                                        df_p_status['validity_clean'] = df_p_status['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                                        df_p_status['route_id_clean'] = df_p_status['route_id'].astype(str).str.strip()
                                        df_p_status['round_clean'] = pd.to_numeric(df_p_status['round'], errors='coerce').fillna(1).astype(int)
                                        
                                        # Bersihkan parameter pembanding dari dashboard
                                        clean_sel_val = str(sel_val).replace(" ", "").lower().strip()
                                        clean_email = str(email).lower().strip()
                                        
                                        
                                        sub_p = df_p_status[
                                            (df_p_status['vendor_email_clean'] == clean_email) & 
                                            (df_p_status['validity_clean'] == clean_sel_val) & 
                                            (df_p_status['route_id_clean'].isin(r_data['route_id'].astype(str).str.strip())) & 
                                            (df_p_status['round_clean'] == int(sel_round))
                                        ]
                                        
                                        if not sub_p.empty:
                                            if "Need Revision" in sub_p['status'].values:
                                                status_ui = '<span class="status-pending" style="color: #9A3412 !important; background-color: #FFEDD5 !important; border-color: #FDBA74 !important;">⚠️ Revisi</span>'
                                            else:
                                                status_ui = '<span class="status-done">✅ Sudah Terisi</span>'
                                                
                                            if "Locked" in sub_p['status'].values: 
                                                is_locked_btn = True
                                    
                                    c1, c2, c3, c4 = st.columns([3, 4, 2, 2])
                                    c1.write(f"**{grp_name}**")
                                    
                                    # PENGAMAN: Cek apakah kolom kota_tujuan benar-benar ada
                                    dests = r_data['kota_tujuan'].unique().tolist() if (not r_data.empty and 'kota_tujuan' in r_data.columns) else []
                                    
                                    if len(dests) > 5: preview_txt = f"{', '.join(dests[:5])}, +{len(dests)-5} kota lainnya"
                                    else: preview_txt = ", ".join(dests)
                                    c2.markdown(f"<span class='route-dest-list'>{preview_txt}</span>", unsafe_allow_html=True)

                                    if is_locked_btn:
                                        c3.button("🔒 Locked", key=f"btn_lk_{gid}_{sel_round}", disabled=True)
                                    else:
                                        if c3.button("📌 Isi Harga", key=f"btn_{t_code}_{gid}_{sel_round}", type="primary"):
                                            st.session_state.update({
                                                'sel_origin': org, 'sel_validity': sel_val, 'sel_load': t_code, 
                                                'vendor_step': 'input', 'focused_group_id': gid, 'sel_round': sel_round
                                            })
                                            st.rerun()
                                    c4.markdown(status_ui, unsafe_allow_html=True)
                                    st.markdown("<hr>", unsafe_allow_html=True)
# --- TAB 3: DOWNLOAD & UPLOAD SPH RESMI ---
        if has_locked:
            t3 = tabs_ui[2]
            with t3:
                st.markdown("### 📄 SPH (Surat Penawaran Harga)")
                st.info("Anda dapat mendownload draf SPH kapan saja. Namun, **fitur Upload baru terbuka setelah SEMUA rute Anda di-Lock oleh Admin**.")
            
                df_p = get_data("Price_Data")
                df_r = get_data("Master_Routes")
                df_g = get_data("Master_Groups")
                df_prof = get_data("Vendor_Profile")
            
                # ▼▼▼ MULAI DARI SINI SUDAH SAYA DORONG KE KANAN (DI DALAM with t3:) ▼▼▼
                if df_p.empty or df_g.empty:
                    st.warning("Belum ada data penawaran.")
                else:
                    # 1. Filter Milik Vendor Ini (SEMUA STATUS: Open & Locked) dan Harga > 0
                    my_prices = df_p[df_p['vendor_email'] == email].copy()
                    my_prices['price'] = pd.to_numeric(my_prices['price'], errors='coerce').fillna(0)
                    my_prices = my_prices[my_prices['price'] > 0]
                    
                    if my_prices.empty:
                        st.warning("Belum ada data harga Anda, atau semua harga Anda masih Rp 0.")
                    else:
                        # 2. Merge Data
                        my_prices['route_id'] = my_prices['route_id'].astype(str).str.strip()
                        df_r['route_id'] = df_r['route_id'].astype(str).str.strip()
                        df_g['group_id'] = df_g['group_id'].astype(str).str.strip()
                        if 'round' not in my_prices.columns: my_prices['round'] = '1'
                        
                        m1 = pd.merge(my_prices, df_r[['route_id', 'group_id', 'kota_tujuan']], on='route_id', how='left')
                        df_final = pd.merge(m1, df_g[['group_id', 'origin', 'load_type']], on='group_id', how='left')
                        
                        # --- TAMBAHAN MERGE MULTIDROP & CATATAN ---
                        df_m = get_data("Multidrop_Data")
                        if not df_m.empty:
                            df_m['group_id'] = df_m['group_id'].astype(str).str.strip()
                            df_m['md_round'] = df_m['id_multidrop'].apply(lambda x: str(x).split('_')[-1] if '_' in str(x) else '1')
                            df_m_sub = df_m[['vendor_email', 'validity', 'group_id', 'md_round', 'inner_city_price', 'outer_city_price', 'labor_cost', 'catatan_tambahan']]
                            
                            df_final = pd.merge(
                                df_final, df_m_sub, 
                                left_on=['vendor_email', 'validity', 'group_id', 'round'],
                                right_on=['vendor_email', 'validity', 'group_id', 'md_round'],
                                how='left'
                            )
                        else:
                            for c in ['inner_city_price', 'outer_city_price', 'labor_cost', 'catatan_tambahan']: df_final[c] = 0
                            
                        # 3. UI Filter (Periode, Load Type, Tahap)
                        c1, c2, c3 = st.columns(3)
                        avail_val = sorted(df_final['validity'].unique().tolist())
                        sel_val = c1.selectbox("Pilih Periode", avail_val, key="sph_val")
                        
                        avail_lt = sorted(df_final['load_type'].dropna().unique().tolist())
                        sel_lt = c2.selectbox("Pilih Tipe Armada", avail_lt, key="sph_lt")
                        
                        avail_rnd = sorted(df_final['round'].unique().tolist())
                        sel_rnd = c3.selectbox("Pilih Tahap Penawaran", avail_rnd, key="sph_rnd")
                        
                        df_print = df_final[(df_final['validity'] == sel_val) & (df_final['load_type'] == sel_lt) & (df_final['round'] == sel_rnd)].copy()
                        v_name = st.session_state['user_info'].get('vendor_name', email)
                        
                        # ▼▼▼ TAMBAHAN: Buang rute yang harganya Rp 0 dari SPH ▼▼▼
                        df_print = df_print[df_print['price'] > 0]
                        
                        # --- CEK APAKAH SEMUA STATUS SUDAH LOCKED ---
                        is_all_locked = False
                        if not df_print.empty:
                            # Cek apakah seluruh baris di tabel df_print statusnya 'Locked'
                            is_all_locked = (df_print['status'] == 'Locked').all()
                        
                        # --- BAGIAN A: DOWNLOAD SPH ---
                        with st.container(border=True):
                            st.markdown("#### Step 1: Download SPH")
                            if df_print.empty:
                                st.info("Tidak ada data untuk kombinasi filter ini.")
                            else:
                                st.write(f"Ditemukan **{len(df_print)} rute** yang siap dicetak SPH-nya. Mohon dapat dicap dan ditandatangani, lalu diupload pada langkah selanjutnya.")
                                
                                # --- MEMBUAT 2 TOMBOL BERSEBELAHAN ---
                                c_btn1, c_btn2 = st.columns(2)
                                
                                # TOMBOL KIRI (WORD SPH RESMI)
                                with c_btn1:
                                    if st.button("📄 Buat Dokumen SPH (Word)", type="primary", use_container_width=True):
                                        tpl_sph = "template_sph.docx"
                                        if not os.path.exists(tpl_sph): 
                                            st.error("Sistem error: Template SPH tidak ditemukan di server. Hubungi Admin!")
                                            st.stop()
                                            
                                        with st.spinner("Merakit Dokumen SPH..."):
                                            try:
                                                v_addr = "-"
                                                if not df_prof.empty:
                                                    vp = df_prof[df_prof['email'] == email]
                                                    if not vp.empty: v_addr = str(vp.iloc[0].get('address', '-'))
                                                
                                                df_print_word = df_print.sort_values(by=['origin', 'kota_tujuan', 'unit_type']).reset_index(drop=True)
                                                file_sph = create_docx_sph(tpl_sph, v_name, v_addr, sel_val, sel_lt, sel_rnd, df_print_word)
                                                
                                                with open(file_sph, "rb") as f:
                                                    st.download_button(
                                                        label="⬇️ Klik di sini untuk Download Word", 
                                                        data=f, 
                                                        file_name=f"SPH_{v_name}_{sel_lt}_Tahap{sel_rnd}.docx", 
                                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                                        type="primary",
                                                        use_container_width=True
                                                    )
                                                os.remove(file_sph)
                                            except Exception as e: st.error(f"Gagal: {e}")
                                
                                # TOMBOL KANAN (EXCEL TABEL SAJA)
                                with c_btn2:
                                    df_excel = df_print.sort_values(by=['origin', 'kota_tujuan', 'unit_type']).reset_index(drop=True)
                                    df_excel['No'] = range(1, len(df_excel) + 1)
                                    
                                    # Rapikan isi kolom sebelum di-export
                                    def fmt_rp(x):
                                        try: return f"Rp {int(x):,}".replace(",", ".")
                                        except: return "Rp 0"
                                    df_excel['Harga Penawaran'] = df_excel['price'].apply(fmt_rp)
                                    df_excel['Multidrop Dalam Kota'] = df_excel['inner_city_price'].apply(fmt_rp)
                                    df_excel['Multidrop Luar Kota'] = df_excel['outer_city_price'].apply(fmt_rp)
                                    df_excel['Biaya Buruh'] = df_excel['labor_cost'].apply(fmt_rp)
                                    
                                    def fmt_lt(x):
                                        x_str = str(x)
                                        if x_str.isdigit() or x_str.replace('.','',1).isdigit(): return f"{x_str} Hari"
                                        return "-" if x_str in ["-", "", "nan", "None"] else x_str
                                    df_excel['Lead Time (Hari)'] = df_excel['lead_time'].apply(fmt_lt)
                                    
                                    # Bersihkan teks kosong (PENGAMAN BARU)
                                    if 'keterangan' not in df_excel.columns: df_excel['keterangan'] = '-'
                                    df_excel['keterangan'] = df_excel['keterangan'].fillna('-').astype(str).replace(['nan', 'None', ''], '-')
                                    
                                    if 'catatan_tambahan' not in df_excel.columns: df_excel['catatan_tambahan'] = '-'
                                    df_excel['catatan_tambahan'] = df_excel['catatan_tambahan'].fillna('-').astype(str).replace(['nan', 'None', ''], '-')
                                    
                                    # Ambil kolom yang diperlukan saja dan ganti judulnya
                                    cols_to_keep = ['No', 'origin', 'kota_tujuan', 'unit_type', 'Lead Time (Hari)', 'Harga Penawaran', 'keterangan', 'Multidrop Dalam Kota', 'Multidrop Luar Kota', 'Biaya Buruh', 'catatan_tambahan']
                                    df_excel = df_excel[cols_to_keep]
                                    df_excel.columns = ['No', 'Origin', 'Tujuan', 'Unit', 'Lead Time', 'Harga Penawaran', 'Keterangan Rute', 'MD Dalam Kota', 'MD Luar Kota', 'Biaya Buruh', 'Catatan Tambahan Vendor']
                                    
                                    # Ubah ke format file Excel (BytesIO)
                                    output = io.BytesIO()
                                    with pd.ExcelWriter(output, engine='openpyxl') as writer:
                                        df_excel.to_excel(writer, index=False, sheet_name='Tabel SPH')
                                    excel_data = output.getvalue()
                                    
                                    safe_ven_xls = "".join(x for x in v_name if x.isalnum())
                                    
                                    # Munculkan tombol Download langsung
                                    st.download_button(
                                        label="📊 Download Tabel (Excel)", 
                                        data=excel_data, 
                                        file_name=f"Tabel_SPH_{safe_ven_xls}_{sel_lt}_Tahap{sel_rnd}.xlsx", 
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                        type="primary",
                                        use_container_width=True
                                    )

                        # --- BAGIAN B: UPLOAD SPH ---
                        with st.container(border=True):
                            st.markdown("#### Step 2: Upload SPH yang Sudah di Cap & Tanda Tangan")
                            
                            if not is_all_locked:
                                # 🛑 Jika belum locked semua, fitur upload disembunyikan dan muncul peringatan
                                st.warning("⏳ **Fitur Upload Terkunci!**\nAdmin TACO harus me-Lock SEMUA penawaran Anda di periode dan tahap ini terlebih dahulu sebelum Anda bisa mengupload dokumen SPH final.")
                            else:
                                # ✅ Jika sudah locked semua, tampilkan uploader
                                st.write(f"Upload untuk: **{sel_lt} | Periode {sel_val} | Tahap {sel_rnd}**")
                                
                                uploaded_file = st.file_uploader("Pilih file SPH (PDF)", type=['pdf', 'png', 'jpg', 'jpeg'])
                                
                                if st.button("📤 Upload Dokumen SPH", type="primary", use_container_width=True):
                                    if uploaded_file is not None:
                                        # MASUKKAN ID FOLDER GOOGLE DRIVE ANDA DI SINI
                                        DRIVE_FOLDER_ID = "0AMIguJ49asOLUk9PVA" 
                                        
                                        safe_ven = "".join(x for x in v_name if x.isalnum())
                                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                        ext = uploaded_file.name.split(".")[-1]
                                        new_filename = f"SPH_{safe_ven}_{sel_lt}_T{sel_rnd}_{timestamp}.{ext}"
                                        
                                        with st.spinner("⏳ Sedang mengirim file ke Google Drive... Mohon tunggu..."):
                                            file_url = upload_to_drive(uploaded_file, new_filename, uploaded_file.type, DRIVE_FOLDER_ID)
                                            
                                            if file_url:
                                                id_up = f"UPL_{safe_ven}_{timestamp}"
                                                save_data("SPH_Uploads", [[id_up, email, v_name, sel_val, sel_lt, sel_rnd, file_url, timestamp]])
                                                st.success("✅ Sukses! Dokumen SPH telah tersimpan aman di Server TACO.")
                                            else:
                                                st.error("⚠️ Gagal mengirim dokumen ke server. Coba lagi.")
                                    else:
                                        st.error("⚠️ Silakan pilih file terlebih dahulu sebelum klik Upload.")

       
    # --- STEP 2: INPUT HARGA ---
    elif step == "input":
        if st.session_state.get('temp_success_msg'):
            st.success(st.session_state['temp_success_msg'])
            st.session_state['temp_success_msg'] = None

        if st.button("⬅️ Kembali ke Menu Utama", type="secondary"):
            st.session_state['vendor_step'] = "dashboard"; st.rerun()

        cur_org = st.session_state.get('sel_origin')
        cur_val = st.session_state.get('sel_validity')
        cur_load = st.session_state.get('sel_load')
        cur_round = str(st.session_state.get('sel_round'))
        focused_gid = st.session_state.get('focused_group_id')

        try: prev_round = str(int(cur_round) - 1)
        except: prev_round = "0"

        st.markdown(f"### Input Penawaran Harga {cur_load}: {cur_org}")
        st.caption(f"Periode: {cur_val} | **Tahap Penawaran: {cur_round}**")

        
        df_acc = get_data("Access_Rights"); df_grps = get_data("Master_Groups")
        if 'round' not in df_acc.columns: df_acc['round'] = '1'
        my_acc = df_acc[(df_acc['vendor_email']==email) & (df_acc['validity']==cur_val) & (df_acc['round']==cur_round)]
        
        target_gids = []
        grp_names = {}
        for gid in my_acc['group_id'].unique():
            r = df_grps[df_grps['group_id']==gid]
            if not r.empty:
                rr = r.iloc[0]
                if rr['origin']==cur_org and rr['load_type']==cur_load:
                    target_gids.append(gid); grp_names[gid]=rr['route_group']
        
        if not target_gids: 
            st.error("Data error."); return

        target_gids = sorted(target_gids)
        if focused_gid and focused_gid in target_gids:
            target_gids.remove(focused_gid)
            target_gids.insert(0, focused_gid)
        
        tabs = st.tabs([grp_names[g] for g in target_gids])
        
        df_r = get_data("Master_Routes"); df_u = get_data("Master_Units"); df_p = get_data("Price_Data"); df_m = get_data("Multidrop_Data")
        if not df_p.empty and 'round' not in df_p.columns: df_p['round'] = '1'
        if not df_m.empty and 'group_id' in df_m.columns: df_m['group_id'] = df_m['group_id'].astype(str).str.strip()

        for i, gid in enumerate(target_gids):
            with tabs[i]:
                g_name = grp_names[gid]
                my_r = df_r[df_r['group_id']==gid]
                my_u = df_u[df_u['group_id']==gid]
                all_u_types = my_u['unit_type'].unique().tolist()
                
                ex_price = {}; ex_spec = {}
                is_lock = False
              
                
                # --- 1. INITIALIZE (Anti-Crash) ---
                source_p_data = pd.DataFrame() 
                current_p_data = pd.DataFrame()
                is_using_prev_data = False
                
                # --- 2. GET CURRENT ROUND DATA ---
                if not df_p.empty:
                    # Bikin teks rujukan dashboard murni rapat & kecil tanpa spasi
                    clean_cur_val = str(cur_val).replace(" ", "").lower().strip()
                    clean_email = str(email).lower().strip()
                    
                    # Bikin dataframe bayangan sementara yang sudah bersih spasi
                    df_p_norm = df_p.copy()
                    df_p_norm['vendor_email_clean'] = df_p_norm['vendor_email'].astype(str).str.strip().str.lower()
                    df_p_norm['validity_clean'] = df_p_norm['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                    df_p_norm['route_id_clean'] = df_p_norm['route_id'].astype(str).str.strip()
                    df_p_norm['round_clean'] = pd.to_numeric(df_p_norm['round'], errors='coerce').fillna(1).astype(int)
                    
                    # Saring data murni berdasarkan format rapat tanpa spasi!
                    current_p_data = df_p_norm[
                        (df_p_norm['vendor_email_clean'] == clean_email) & 
                        (df_p_norm['validity_clean'] == clean_cur_val) & 
                        (df_p_norm['route_id_clean'].isin(my_r['route_id'].astype(str).str.strip())) & 
                        (df_p_norm['round_clean'] == int(cur_round))
                    ]

                
                # --- 3. DETERMINE REFERENCE DATA ---
                # ▼ POTONGAN FIX SAKLEK: SINKRONISASI STRUKTUR REFERENSI BIAR GA CRASH & UNIT MUNCUL ▼
                if current_p_data.empty and str(cur_round) != "1":
                    if not df_p.empty:
                        source_p_data = df_p_norm[
                            (df_p_norm['vendor_email_clean'] == clean_email) & 
                            (df_p_norm['validity_clean'] == clean_cur_val) & 
                            (df_p_norm['route_id_clean'].isin(my_r['route_id'].astype(str).str.strip())) & 
                            (df_p_norm['round_clean'] == int(prev_round))
                        ].copy()
                        is_using_prev_data = True 
                else:
                    # Pastikan jika mengambil dari data ronde sekarang, struktur kolom clean-nya tetep dibawa
                    source_p_data = current_p_data.copy()

                # Pastikan kolom penampung harga dibaca sebagai angka numeric murni biar gak katarak
                if not source_p_data.empty:
                    source_p_data['price_numeric'] = pd.to_numeric(source_p_data['price'], errors='coerce').fillna(0)

                u_types = []
                if cur_round == "1":
                    # Kalau Tahap 1, munculkan semua unit bawaan master groups tanpa saringan
                    u_types = all_u_types
                else:
                    # Kalau Tahap 2 ke atas, ambil unit rujukan ronde lalu yang harganya valid diisi vendor
                    if not source_p_data.empty:
                        submitted_units = source_p_data[source_p_data['price_numeric'] > 0]['unit_type'].unique().tolist()
                        u_types = [u for u in all_u_types if u in submitted_units]
                    
                    # PENGAMAN NUKLIR: Jika vendor belum isi / submitted_units kosong, backup balik ke semua unit master
                    if not u_types:
                        u_types = all_u_types

                # Jika benar-benar tidak ada unit yang terdaftar di master groups, baru kita skip form ini
                if not u_types:
                    st.warning(f"⚠️ Tidak ada jenis unit armada yang terdaftar untuk grup rute {g_name}.")
                    continue

                # --- STEP 5: POPULATE ex_price & ex_spec ---
                if not source_p_data.empty:
                    if not is_using_prev_data and "Locked" in source_p_data['status'].values:
                        is_lock = True
                    for _, row in source_p_data.iterrows():
                        harga_bersih = clean_numeric(row['price'])
                        rid_clean = str(row['route_id']).strip()
                        ex_price[(rid_clean, row['unit_type'])] = int(harga_bersih) if harga_bersih else 0
                        ex_spec[row['unit_type']] = {'w': row.get('weight_capacity'), 'c': row.get('cubic_capacity')}

                with st.form(key=f"f_{gid}_{cur_round}"):
                    # 1. SPEC
                    with st.container(border=True):
                        st.markdown(f"#### 🛻 Spesifikasi Armada")
                        if is_using_prev_data: st.info(f"ℹ️ Data disalin dari Tahap {prev_round}.")
                        
                        sp_data = []
                        for u in u_types:
                            sp_data.append({
                                "Jenis Unit": u,
                                "Kapasitas Berat Bersih (Kg)": clean_numeric(ex_spec.get(u,{}).get('w')),
                                "Kapasitas Kubikasi Dalam (CBM)": clean_numeric(ex_spec.get(u,{}).get('c'))
                            })
                        
                        df_sp = pd.DataFrame(sp_data)
                        cf_sp = {
                            "Jenis Unit": st.column_config.TextColumn(disabled=True),
                            "Kapasitas Berat Bersih (Kg)": st.column_config.NumberColumn(min_value=0, format="%,d", step=1),
                            "Kapasitas Kubikasi Dalam (CBM)": st.column_config.NumberColumn(min_value=0, format="%.2f", step=0.1)
                        }
                        ed_sp = st.data_editor(df_sp, hide_index=True, use_container_width=True, disabled=is_lock, column_config=cf_sp)

                    # 2. PRICE
                    with st.container(border=True):
                        st.markdown(f"#### 💰 Penawaran Harga")
                        if cur_round == "2":
                            st.info("💡 **Catatan untuk Tahap 2:**\n\n"
                                "🎯 = **Target Price** (Harga referensi dari tahap sebelumnya).\n\n"
                                "Mohon dapat mengganti harga penawaran di kolom harga **mendekati atau lebih kompetitif** dari target price yang tertera sebagai penawaran harga Tahap 2.")
                        p_data = []
                        
                        # --- 1. SUSUN DATA PER BARIS ---
                        for _, row in my_r.iterrows():
                            rid = str(row['route_id']).strip()
                            rd = {
                                "Route ID": rid, 
                                "Kota Asal": row['kota_asal'], 
                                "Kota Tujuan": row['kota_tujuan'],
                                "Keterangan": row.get('keterangan', '-'),
                                "Lead Time (Hari)": 0 
                            }
                            
                            # Ambil Lead Time histori
                            # ▼ TIMPA TIGA BARIS TADI PAKAI BLOK INI SAKLEK ▼
                            if not source_p_data.empty:
                                match_rows = source_p_data[source_p_data['route_id'].astype(str).str.strip() == str(rid).strip()]
                                if not match_rows.empty:
                                    val_lt_raw = match_rows.iloc[0].get('lead_time', 0)
                                    cleaned_lt = clean_numeric(val_lt_raw)
                                    if cleaned_lt > 0:
                                        rd["Lead Time (Hari)"] = int(cleaned_lt)
                                    else:
                                        try: rd["Lead Time (Hari)"] = int(float(str(val_lt_raw).strip()))
                                        except: rd["Lead Time (Hari)"] = 0

                            # Input Harga & Target per Unit
                            for u in u_types:
                                # Kolom Target (Hanya muncul di Round 2)
                                if cur_round == "2":
                                    tgt = get_target_price(df_p, rid, u, cur_val)
                                    rd[f"Target {u}"] = f"Rp {int(tgt):,}".replace(",", ".") if tgt > 0 else "-"
                                
                                # Kolom Input Harga
                                rd[f"Harga {u} per trip"] = ex_price.get((rid, u), 0)
                            
                            p_data.append(rd)

                        # --- 2. SETUP DATAFRAME & URUTAN KOLOM ---
                        if not p_data:
                            st.warning("⚠️ Rute tidak ditemukan.")
                            df_pr = pd.DataFrame()
                        else:
                            df_pr = pd.DataFrame(p_data)
                            
                            # Susun urutan kolom agar rapi
                            cols_order = ["Route ID", "Kota Asal", "Kota Tujuan", "Lead Time (Hari)"]
                            for u in u_types:
                                if cur_round == "2": cols_order.append(f"Target {u}")
                                cols_order.append(f"Harga {u} per trip")
                            cols_order.append("Keterangan")
                            
                            # Filter hanya kolom yang beneran ada di df_pr
                            final_cols = [c for c in cols_order if c in df_pr.columns]
                            df_pr = df_pr[final_cols]

                        # --- 3. CONFIGURASI KOLOM (Kunci Target, Buka Harga) ---
                        cf_pr = {
                            "Route ID": None,
                            "Kota Asal": st.column_config.TextColumn(disabled=True),
                            "Kota Tujuan": st.column_config.TextColumn(disabled=True),
                            "Lead Time (Hari)": st.column_config.NumberColumn(format="%d", min_value=0),
                            "Keterangan": st.column_config.TextColumn(),
                        }
                        
                        for u in u_types:
                            # Harga BISA EDIT
                            cf_pr[f"Harga {u} per trip"] = st.column_config.NumberColumn(
                                label=f"💰Harga {u} per trip", min_value=0, step=1000, format="Rp %,d"
                            )
                            # Target MATI (Locked)
                            if f"Target {u}" in df_pr.columns:
                                cf_pr[f"Target {u}"] = st.column_config.TextColumn(
                                    label=f"🎯 Target {u}", disabled=True, width=90, help="Target Price"
                                )

                        # --- 4. TAMPILKAN EDITOR (Tanpa .style agar bisa diedit) ---
                        ed_pr = st.data_editor(
                            df_pr, 
                            hide_index=True, 
                            use_container_width=True, 
                            column_config=cf_pr,
                            key=f"editor_state_{gid}_{cur_round}",
                            disabled=is_lock
                        )
                    
                    # 3. MULTIDROP
                    with st.container(border=True):
                        st.markdown("#### 📦 Biaya Multidrop & Buruh")
                        ic, oc, lc = 0, 0, 0
                        md_source = pd.DataFrame()
                        
                        # ▼ POTONGAN FIX SAKLEK: BIAR HARGA MULTIDROP LAMA KELOAD SEMPURNA DI FORM VENDOR ▼
                        if not df_m.empty:
                            # Bikin dataframe bayangan khusus multidrop yang bersih spasi
                            df_m_norm = df_m.copy()
                            df_m_norm['vendor_email_clean'] = df_m_norm['vendor_email'].astype(str).str.strip().str.lower()
                            df_m_norm['validity_clean'] = df_m_norm['validity'].astype(str).str.replace(" ", "").str.lower().str.strip()
                            df_m_norm['group_id_clean'] = df_m_norm['group_id'].astype(str).str.strip()
                            
                            # Bersihkan pembanding dari dashboard
                            clean_cur_val = str(cur_val).replace(" ", "").lower().strip()
                            clean_email = str(email).lower().strip()
                            
                            # Saring data murni tanpa sensitif spasi
                            base_m = df_m_norm[
                                (df_m_norm['vendor_email_clean'] == clean_email) & 
                                (df_m_norm['validity_clean'] == clean_cur_val) & 
                                (df_m_norm['group_id_clean'] == str(gid).strip())
                            ]
                            
                            if not base_m.empty:
                                # Cari data ronde berjalan atau ronde sebelumnya lewat id_multidrop
                                md_curr = base_m[base_m['id_multidrop'].astype(str).str.endswith(f"_{cur_round}")]
                                if not md_curr.empty: 
                                    md_source = md_curr
                                elif str(cur_round) != "1":
                                    md_prev = base_m[base_m['id_multidrop'].astype(str).str.endswith(f"_{prev_round}")]
                                    if not md_prev.empty: md_source = md_prev
                                    
                        if not md_source.empty:
                            ic = clean_numeric(md_source.iloc[0].get('inner_city_price')) or 0
                            oc = clean_numeric(md_source.iloc[0].get('outer_city_price')) or 0
                            lc = clean_numeric(md_source.iloc[0].get('labor_cost')) or 0
                        
                        df_md_ui = pd.DataFrame([{"Multidrop Dalam Kota": ic, "Multidrop Luar Kota": oc, "Biaya Buruh": lc}])
                        
                        cf_md = {
                            "Multidrop Dalam Kota": st.column_config.NumberColumn(min_value=0, step=1000, format="Rp %,d"),
                            "Multidrop Luar Kota": st.column_config.NumberColumn(min_value=0, step=1000, format="Rp %,d"),
                            "Biaya Buruh": st.column_config.NumberColumn(min_value=0, step=1000, format="Rp %,d")
                        }
                        ed_md = st.data_editor(df_md_ui, hide_index=True, use_container_width=True, disabled=is_lock, column_config=cf_md, key=f"ed_md_{gid}_{cur_round}")
                        st.markdown("<br><b>📝 Catatan Tambahan (Opsional)</b>", unsafe_allow_html=True)
                        
                        prev_note = ""
                        if not md_source.empty and 'catatan_tambahan' in md_source.columns:
                            prev_note = str(md_source.iloc[0]['catatan_tambahan'])
                            if prev_note.lower() == "nan": prev_note = ""
                            
                        vendor_note = st.text_area(
                            "Keterangan lain yang perlu diinfokan (jika ada):", 
                            value=prev_note, 
                            disabled=is_lock,
                            height=100
                        )
                    # SAVE BUTTON
                    st.write("")
                    if st.form_submit_button(f"Simpan Data {cur_load} {g_name} (Tahap {cur_round})", type="primary") and not is_lock:
                        c_spec = {r['Jenis Unit']: {'w': r['Kapasitas Berat Bersih (Kg)'], 'c': r['Kapasitas Kubikasi Dalam (CBM)']} for _, r in ed_sp.iterrows()}
                        
                        f_data = []
                        ts = (datetime.utcnow() + timedelta(hours=7)).strftime("%Y-%m-%d %H:%M:%S")
                        editor_key = f"editor_state_{gid}_{cur_round}"
                        if editor_key in st.session_state and st.session_state[editor_key]:
                            # Streamlit menyimpan perubahan di session state, kita timpa ke df_pr asli
                            changes = st.session_state[editor_key]
                            df_terupdate = df_pr.copy()
                            
                            # Aplikasikan perubahan sel yang diedit vendor secara real-time
                            if "edited_rows" in changes:
                                for row_idx, changed_cols in changes["edited_rows"].items():
                                    for col_name, new_val in changed_cols.items():
                                        df_terupdate.at[int(row_idx), col_name] = new_val
                        else:
                            df_terupdate = df_pr
                            
                        for _, r in df_terupdate.iterrows():
                            rid = str(r['Route ID']).replace(" ", "").strip()
                            lt = int(clean_numeric(r['Lead Time (Hari)']) or 0)
                            ket = str(r['Keterangan']).replace("nan", "-")
                            
                            for u in u_types:
                                pr = int(clean_numeric(r[f"Harga {u} per trip"]) or 0)
                                w = str(c_spec.get(u,{}).get('w','')).replace("nan", "-"); c = str(c_spec.get(u,{}).get('c','')).replace("nan", "-")
                                round_num = int(cur_round)
                                tid = f"{email}_{cur_val}_{rid}_{u}_{round_num}".replace(" ","")
                                if pr > 0:
                                    f_data.append([tid, email, "Open", cur_val, rid, u, lt, pr, w, c, ket, ts, round_num])
                        
                        # ▼ POTONGAN 3: TIMPA 4 BARIS LAMA PAKAI INI BIAR INPUTAN KEBACA ▼
                        df_md_terupdate = df_md_ui.copy()
                        editor_md_dyn_key = f"ed_md_{gid}_{cur_round}"
                        
                        if editor_md_dyn_key in st.session_state and st.session_state[editor_md_dyn_key]:
                            changes_md = st.session_state[editor_md_dyn_key]
                            if "edited_rows" in changes_md:
                                for row_idx, changed_cols in changes_md["edited_rows"].items():
                                    for col_name, new_val in changed_cols.items():
                                        df_md_terupdate.at[int(row_idx), col_name] = new_val
                        
                        try: mi = int(float(str(df_md_terupdate.iloc[0]["Multidrop Dalam Kota"]).replace("Rp","").replace(" ","").strip()))
                        except: mi = 0
                        
                        try: mo = int(float(str(df_md_terupdate.iloc[0]["Multidrop Luar Kota"]).replace("Rp","").replace(" ","").strip()))
                        except: mo = 0
                        
                        try: ml = int(float(str(df_md_terupdate.iloc[0]["Biaya Buruh"]).replace("Rp","").replace(" ","").strip()))
                        except: ml = 0
                        
                        val_no_space = str(cur_val).replace(" ", "").strip()
                        mid = f"M_{email}_{gid}_{val_no_space}_{int(cur_round)}"

                        # ▼▼▼ PERBAIKAN: CEK STATUS KEBERHASILAN SAVE ▼▼▼
                        with st.spinner("Menyimpan ke server..."):
                            res_price = save_data("Price_Data", f_data)
                            res_md = save_data("Multidrop_Data", [[mid, email, cur_val, gid, mi, mo, ml, ts, vendor_note]])
                            
                            if res_price and res_md:
                                st.session_state['temp_success_msg'] = f"Sukses! Data Tahap {cur_round} tersimpan."
                                st.cache_data.clear()
                                st.rerun()
                            else:
                                st.error("⚠️ Proses simpan terganggu oleh server Google. Silakan klik Simpan 1x lagi.")
                        
                        
if __name__ == "__main__":
    main()




